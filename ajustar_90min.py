"""
Ajusta la presentación a 90 minutos:
1. Marca slides opcionales con data-optional="true" en el HTML
2. Agrega CSS/JS para mostrar indicador "OPCIONAL · Saltar →" en pantalla
3. Regenera el Word con el guión de 90 min
"""
import re, pathlib
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────
# 1. MODIFICAR HTML
# ─────────────────────────────────────────────────────────────
html_path = pathlib.Path(r"C:\Users\HP\Documents\mleyvaz.github.io\clase_espejo_tercera_respuesta.html")
html = html_path.read_text(encoding="utf-8")

# Slides a marcar como opcionales (identificados por su comentario único)
OPTIONAL_COMMENTS = [
    "SLIDE 3-NEW: WARM-UP POLL",
    "SLIDE 6: EXPERIMENT 1 — HALLUCINATION ATTACK VECTORS",
    "SLIDE 6B: VECTOR B",
    "SLIDE 6C: VECTOR C",
    "SLIDE 9: FAMILY OF UNCERTAINTY FRAMEWORKS",
    "SLIDE NEW: ALEATORIC vs EPISTEMIC UNCERTAINTY",
    "SLIDE NEW: NEUTROSOPHIC AS ONE TOOL",
    "SLIDE NEW: AI QUERY LAB",
    "SLIDE 14 (was 14): EXPERIMENT 2 — TRUST CALIBRATION",
    "SLIDE NEW: SELF-CALIBRATION EXERCISE",
    "SLIDE 15 (was 15): PARACONSISTENCY",
    "SLIDE 16: ABSTENTION",
    "SLIDE 17: APPLIED CASE 2 — FRAUD",
    "SLIDE 18: TRAJECTORY (COMPASS TO GPS)",
    "SLIDE 19: LIVE POLL",
    "SLIDE NEW: CONFORMAL PREDICTION",
    "SLIDE NEW: CAUSALITY APPLIED",
    "SLIDE 23: FROM DASHBOARD TO COMPASS",
    "SLIDE 24-NEW: INTERACTIVE BI DASHBOARD",
    "SLIDE 25 (was 24): WHAT TO BUILD",
    "SLIDE 25: WHAT THIS IS NOT",
    "SLIDE 26: PRACTICAL TOOLKIT",
    "SLIDE NEW: OOD DETECTION",
]

marked = 0
for comment_text in OPTIONAL_COMMENTS:
    # Match the comment line + optional whitespace + the opening section tag
    pattern = (
        r'(<!--\s*={3,}\s*' + re.escape(comment_text) + r'.*?={3,}\s*-->)'
        r'(\s*<section class="slide(?:\s+light)?")([^>]*>)'
    )
    def add_optional(m):
        # Avoid double-marking
        if 'data-optional' in m.group(3):
            return m.group(0)
        return m.group(1) + m.group(2) + ' data-optional="true"' + m.group(3)[-1:]

    new_html, n = re.subn(pattern, add_optional, html, count=1)
    if n:
        html = new_html
        marked += 1
    else:
        print(f"  [!] No match: {comment_text}")

print(f"Slides marcadas como opcionales: {marked}/{len(OPTIONAL_COMMENTS)}")

# ── CSS para el indicador flotante ──
CSS_INJECT = """
  /* ============ OPTIONAL SLIDE INDICATOR ============ */
  #optionalBadge {
    position: fixed;
    top: 16px;
    right: 20px;
    z-index: 500;
    display: none;
    align-items: center;
    gap: 10px;
    background: rgba(240,160,32,0.15);
    border: 1.5px solid rgba(240,160,32,0.55);
    border-radius: 100px;
    padding: 7px 16px 7px 12px;
    font-family: 'JetBrains Mono', monospace;
    font-size: 10px;
    letter-spacing: 0.12em;
    color: #f0a020;
    cursor: pointer;
    transition: opacity 0.3s;
    user-select: none;
  }
  #optionalBadge.visible { display: flex; }
  #optionalBadge:hover { background: rgba(240,160,32,0.25); }
"""

# ── JS para el indicador flotante ──
JS_INJECT = """
  // ===== OPTIONAL SLIDE INDICATOR =====
  const optBadge = document.createElement('div');
  optBadge.id = 'optionalBadge';
  optBadge.innerHTML = '⏭ OPCIONAL &nbsp;·&nbsp; Saltar →';
  optBadge.title = 'Este slide no es necesario en la versión de 90 min. Pulsa para saltar.';
  document.body.appendChild(optBadge);

  function isOptional(idx) {
    return slides[idx] && slides[idx].dataset.optional === 'true';
  }

  function skipToNext() {
    let next = current + 1;
    while (next < total && isOptional(next)) next++;
    if (next < total) goTo(next);
  }

  optBadge.addEventListener('click', skipToNext);

  // Hook into updateUI
  const _origUpdateUI = updateUI;
  // Override (patch after definition) — done inline below via monkey-patch
"""

# Inyectar CSS antes de </style>
html = html.replace("</style>", CSS_INJECT + "\n</style>", 1)

# Inyectar el badge div y lógica JS justo antes del cierre </script> del bloque principal
# Encontramos el bloque que contiene updateUI
insert_marker = "  updateUI();\n\n  // ===== INTERACTIVE COMPASS ====="
js_patch = """  updateUI();

  // ===== OPTIONAL BADGE (90-min mode) =====
  (function() {
    const optBadge = document.createElement('div');
    optBadge.id = 'optionalBadge';
    optBadge.innerHTML = '⏭ OPCIONAL &nbsp;·&nbsp; <strong>Saltar →</strong>';
    optBadge.title = 'No es necesario en 90 min. Clic para saltar.';
    document.body.appendChild(optBadge);

    function refreshBadge() {
      if (slides[current] && slides[current].dataset.optional === 'true') {
        optBadge.classList.add('visible');
      } else {
        optBadge.classList.remove('visible');
      }
    }

    optBadge.addEventListener('click', () => {
      let next = current + 1;
      while (next < total && slides[next] && slides[next].dataset.optional === 'true') next++;
      if (next < total) goTo(next);
    });

    // Patch updateUI to also refresh badge
    const _orig = updateUI;
    window.updateUI = function() { _orig(); refreshBadge(); };
    // Override references
    const origGoTo = goTo;
    // Re-wire: call refreshBadge after any navigation
    document.addEventListener('keydown', refreshBadge, true);
    document.getElementById('prevBtn').addEventListener('click', refreshBadge);
    document.getElementById('nextBtn').addEventListener('click', refreshBadge);
    refreshBadge();
  })();

  // ===== INTERACTIVE COMPASS ====="""

html = html.replace(insert_marker, js_patch, 1)

html_path.write_text(html, encoding="utf-8")
print("HTML guardado con indicador de slides opcionales.")

# ─────────────────────────────────────────────────────────────
# 2. GENERAR WORD — GUIÓN 90 MIN
# ─────────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

def slide_label(doc, text, skip=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    fill = "6B4200" if skip else "0A3A52"
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)
    return p

def time_tag(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text); run.italic = True; run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

def body(doc, text):
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Cm(0.4)
    p.add_run(text).font.size = Pt(11)

def tip(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("💡 NOTA: " + text)
    r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xC4, 0x7D, 0x10)

def interaction(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("🎯 INTERACCIÓN: " + text)
    r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x00, 0x96, 0xB8)

def separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    r = p.add_run("─" * 60); r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# ─── TÍTULO ───
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_after = Pt(4)
r = tp.add_run("GUIÓN 90 MIN — La tercera respuesta en analítica de datos")
r.bold = True; r.font.size = Pt(17); r.font.color.rgb = RGBColor(0x0A, 0x16, 0x28)

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER; sp.paragraph_format.space_after = Pt(2)
sp.add_run("UNIR · UBE · 25 de mayo de 2026 · 20:00  |  Dr. Maikel Leyva-Vázquez").font.size = Pt(10)

note_p = doc.add_paragraph()
note_p.paragraph_format.space_before = Pt(8); note_p.paragraph_format.space_after = Pt(12)
note_p.paragraph_format.left_indent = Cm(0.4)
nr = note_p.add_run(
    "Las secciones con fondo NARANJA son opcionales: la presentación las marca con un botón "
    "\"⏭ OPCIONAL · Saltar →\" visible solo para el presentador. Sáltalas si el tiempo aprieta. "
    "Las frases entre comillas son literales sugeridos. Tiempo total estimado: 85–92 min."
)
nr.italic = True; nr.font.size = Pt(10); nr.font.color.rgb = RGBColor(0x5A, 0x6C, 0x85)
separator(doc)

# ─── ESTRUCTURA VISUAL DE 90 MIN ───
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
r = p.add_run("MAPA DE TIEMPO — 90 MIN"); r.bold = True; r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x0A, 0x16, 0x28)

bloques = [
    ("ACTO I — El problema · ~22 min",
     [("Cover + Apertura + Agenda + LinkedIn", "4 min"),
      ("Hook + Schwartz + Paralelo analítica", "6 min"),
      ("Experimento VECTOR A (único activo)", "6 min"),
      ("Arquitectura + La tercera respuesta + GPS/Brújula", "4 min"),
      ("LR1: LLMs razonan por patrón, no por prueba", "2 min")]),
    ("ACTO II — El marco · ~27 min",
     [("LR2: 3 técnicas logicas — CoT + ReAct + ToT + plantillas", "5 min"),
      ("Las cuatro zonas (interacción: paracetamol headline)", "5 min"),
      ("Brújula interactiva (tu decisión de este mes)", "5 min"),
      ("Caso churn", "3 min"),
      ("Plantilla prompt T-I-F = CoT epistémico", "3 min"),
      ("Escalera de Pearl — causalidad", "4 min"),
      ("Ejemplo tutoría (verbalmente, sin slide)", "2 min")]),
    ("ACTO III — Aplicación · ~20 min",
     [("Caso autismo / paracetamol prenatal", "5 min"),
      ("Aula Honesta — todas las disciplinas", "3 min"),
      ("Big Idea Recap + 4 lecciones filosóficas", "4 min"),
      ("Pregunta de cierre — 2 min de silencio real", "3 min"),
      ("Encuesta final antes vs después", "5 min")]),
    ("CIERRE · ~25 min",
     [("Q&A abierto", "18 min"),
      ("Speaker card + Créditos", "2 min")]),
]

for bloque_title, items in bloques:
    bp = doc.add_paragraph()
    bp.paragraph_format.space_before = Pt(10); bp.paragraph_format.space_after = Pt(2)
    br = bp.add_run(bloque_title); br.bold = True; br.font.size = Pt(11)
    br.font.color.rgb = RGBColor(0x00, 0x72, 0x94)
    for item, t in items:
        ip = doc.add_paragraph()
        ip.paragraph_format.left_indent = Cm(0.6); ip.paragraph_format.space_after = Pt(1)
        ip.add_run(f"• {item}").font.size = Pt(10)
        tp2 = ip.add_run(f"  [{t}]")
        tp2.font.size = Pt(9); tp2.font.color.rgb = RGBColor(0x88, 0x88, 0x88); tp2.italic = True

separator(doc)

# ─────────────────────────────────────────────────────────────
# GUIÓN POR SLIDE
# ─────────────────────────────────────────────────────────────
KEEP  = False
SKIP  = True

slides_guion = [

("SLIDE 1 — PORTADA", "1 min", KEEP, [
    ("t", "Mientras entran, deja la portada visible. Cuando todos estén:"),
    ("t", '"Buenas noches. Bienvenidos a esta clase espejo entre UNIR y UBE. Me llamo Maikel Leyva-Vázquez, coordino el posgrado en UBE Ecuador y esta noche vamos a hablar de algo que uso todos los días en mi trabajo: cómo enseñarle a la IA a decir \'no lo sé\'."'),
    ("t", '"La clase se llama La tercera respuesta. Ya verán por qué."'),
    ("n", "Pausa 3 segundos. Deja que el título respire."),
]),

("SLIDE 2b — PREGUNTA DE APERTURA", "2 min", KEEP, [
    ("t", '"Antes de abrir cualquier diapositiva, quiero que piensen 30 segundos en alguien a quien respetan profundamente — un mentor, un investigador, un médico, quien sea."'),
    ("t", '"¿Alguna vez los escucharon decir: no sé, necesito más datos?"'),
    ("i", "Silencio real de 10 segundos. Espera que algunos asientan o escriban en el chat."),
    ("t", '"Esa honestidad intelectual es exactamente lo que vamos a construir con datos esta noche — y lo vamos a enseñarle también a la IA."'),
]),

("SLIDE 2 — AGENDA", "1 min", KEEP, [
    ("t", '"Esta noche tiene cinco bloques y ocho actividades prácticas. La meta no es cubrir todo — es que salgan con algo que puedan usar mañana."'),
    ("t", '"Lean el tagline: Esta noche aprenderán a decir no sé de una forma que vale más que cualquier respuesta de ChatGPT. Lo digo en serio."'),
]),

("SLIDE 3-NEW — TERMÓMETRO DE CONFIANZA", "— SALTAR en 90 min —", SKIP, [
    ("n", "SALTAR. Si sobra tiempo al final, puedes hacerlo como actividad de cierre."),
]),

("SLIDE 3c — LINKEDIN LABOR MARKET REPORT", "2 min", KEEP, [
    ("t", '"Ancla en el mundo real: informe LinkedIn enero 2026, mil millones de usuarios."'),
    ("t", '"Solo el 3% declara habilidades de IA — y los Forward-Deployed Engineers crecieron 42 veces en un año. El cuello de botella no es técnica. Es epistémico: ¿quién sabe qué preguntarle a la IA? ¿Cuándo no creerle? Eso es lo de esta noche."'),
    ("i", "Chat rápido: '¿Alguien usa IA en su trabajo diario? Sí / No.'"),
]),

("SLIDE 3 — HOOK", "1 min", KEEP, [
    ("t", '"Tu modelo de IA no te miente. Pero tampoco sabe cuándo está equivocado. Y eso — para quien toma decisiones con datos — es peor que mentir."'),
]),

("SLIDE 4 — CASO SCHWARTZ", "4 min", KEEP, [
    ("t", '"Mayo 2023. Abogado en Nueva York necesitaba jurisprudencia. Le preguntó a ChatGPT. El modelo devolvió seis precedentes perfectamente formateados: número de caso, cita, razonamiento. Todo consistente. Todo inexistente."'),
    ("t", '"Schwartz los incluyó en su escrito. El juez los buscó. No existían. Le preguntó a ChatGPT si eran reales — el modelo confirmó dos veces. Le pidió el texto completo — el modelo lo generó con juez ficticio, citas ficticias, razonamiento ficticio."'),
    ("t", '"Resultado: multa de cinco mil dólares, carrera dañada."'),
    ("t", '"El error no fue confiar en la IA. Fue asumir que su confianza indicaba capacidad de saber cuándo no saber."'),
    ("n", "Pausa dramática. Este caso impacta a todo el mundo."),
]),

("SLIDE 5 — PARALELO ANALÍTICA", "1 min", KEEP, [
    ("t", '"Schwartz pidió precedentes legales. Ustedes le piden predicciones, scoring, recomendaciones clínicas. El mismo sistema, con el mismo 95% de confianza — sin importar si sabe o adivina."'),
]),

("SLIDE 6 — EXPERIMENTO INTRO", "— SALTAR en 90 min —", SKIP, [
    ("n", "SALTAR. Ir directamente a SLIDE 6A."),
]),

("SLIDE 6A — VECTOR A: RESUME EL PAPER", "6 min", KEEP, [
    ("t", '"Vamos a hacer el experimento. Copia este prompt exactamente y pégalo en ChatGPT o Gemini. No cambies ni una coma. Tienes dos minutos."'),
    ("i", "Espera 2 minutos reales. Pide que compartan resultados en el chat. Lee 2-3 en voz alta — probablemente incluyen F1-scores, tamaños de muestra, comparaciones con Random Forest."),
    ("t", '"Ahora búsquenlo. scholar.google.com — ese DOI. ¿Qué encuentran?"'),
    ("t", '"Cero resultados. El paper no existe. Pero el modelo les dio métricas específicas, baselines, conclusiones — todo estadísticamente plausible. No inventó al azar: predijo qué datos debería tener un paper de ese diseño."'),
    ("n", "Si alguien dice que el modelo no hallucinó: felicítalos — usaron un modelo mejor calibrado o reformularon sin querer. Eso también es una lección."),
]),

("SLIDE 6B — VECTOR B", "— SALTAR en 90 min —", SKIP, [
    ("n", "SALTAR en versión 90 min. Disponible para explorar después de clase."),
]),

("SLIDE 6C — VECTOR C", "— SALTAR en 90 min —", SKIP, [
    ("n", "SALTAR en versión 90 min. Disponible para explorar después de clase."),
]),

("SLIDE 7 — PROBLEMA DE ARQUITECTURA", "3 min", KEEP, [
    ("t", '"Ahora entendemos por qué pasa. Un LLM no busca verdad — predice el siguiente token más probable. No hay módulo que distinga \'esto lo sé\' de \'esto suena bien, lo genero\'. La alucinación no es un bug: es el sistema funcionando como fue diseñado."'),
    ("t", '"En nuestro estudio Breaking the Chains — Leyva-Vázquez y Smarandache, NSS 2026 — analizamos 100 consultas a cuatro modelos GPT. El 66% mostró hiper-verdad: T+I+F > 1. El modelo asignaba más certeza total de la que el marco permite. Y aparece en GPT-4, Claude Opus, Gemini — los mejores modelos."'),
]),

("SLIDE 7c — BREAKING THE CHAINS: RESULTADOS COMPLETOS", "3 min", KEEP, [
    ("t", '"Este resultado viene de un paper publicado — nuestro propio estudio. N=100 evaluaciones neutrosóficas válidas, 4 modelos GPT, 5 fenómenos lingüísticos."'),
    ("t", '"El 66% de las evaluaciones sin restricción producen hiper-verdad — T+I+F mayor que 1. IC 95% entre 56 y 74 por ciento. El fenómeno no es ruido."'),
    ("t", '"El más extremo: contradicción ética al 95%. Contingencia futura al 70%. Vaguedad al 60%. Ignorancia epistémica al 55%. Paradoja lógica al 50. Y el prompting probabilístico suprime precisamente esos componentes — Δ_I de +0.383 en ignorancia epistémica."'),
    ("t", '"Mason 2026 replicó independientemente con 5 vendors distintos: Anthropic, Meta, DeepSeek, Alibaba, Mistral. 84% de hiper-verdad. No es un artefacto de OpenAI."'),
    ("n", "Citar: Leyva-Vázquez & Smarandache. NSS Vol. 99, 2026, pp. 288-299. Código: github.com/mleyvaz/neutrosophic-llm-logic. DOI: 10.5281/zenodo.19911845."),
    ("i", "Si alguien en el chat pide el paper: el slide tiene el DOI completo — pueden buscarlo ahora mismo."),
]),

("SLIDE 7b — N-DEL: RESULTADOS EMPÍRICOS", "2 min", KEEP, [
    ("t", '"¿Funciona medir incertidumbre así? Sí — y tenemos los números."'),
    ("t", '"N-DEL — Neutrosophic Dynamic Epistemic Logic — evaluado en 120 preguntas difíciles, dos modelos grandes, 40 anotadores humanos. AUROC 0.919 contra Semantic Entropy 0.47 — el doble de discriminación. 93.5% de abstención cuando el modelo debería abstenerse. Correlación de 0.861 con el juicio humano. La distancia TIF bajó 57% respecto al baseline."'),
    ("t", '"Y más importante: en salud 0.918, en política 0.880, en tecnología emergente 0.887. No es un truco de laboratorio — funciona en las áreas donde más importa."'),
    ("n", "Citar: Leyva-Vázquez & Smarandache (2025). N-DEL. N=120 · 2 modelos · 40 anotadores."),
]),

("SLIDE 7d — VOCES A TRAVÉS DEL TIEMPO", "2 min", KEEP, [
    ("t", '"Antes de que existiera un solo modelo de IA, la humanidad ya sabía que la realidad no cabe en verdadero o falso."'),
    ("t", '"Heráclito, cinco siglos antes de Cristo: \'lo que es opuesto está en concierto — de lo que difiere surge la armonía más bella.\' El Popol Vuh maya: los Creadores deliberan en silencio antes de hablar — no actúan desde la certeza, actúan desde la meditación sobre la incertidumbre. Leibniz sueña con la máquina de razonar — y ya intuye que necesita tres valores, no dos. Bohr lo dice en física cuántica: el opuesto de una verdad profunda puede ser otra verdad profunda."'),
    ("t", '"Lo que Smarandache formalizó en 1995 y lo que nosotros aplicamos a los LLMs esta noche no es un invento — es una herencia de 2.500 años de inteligencia que se negó a simplificar."'),
    ("n", "Pausa de 3 segundos. Deja que las imágenes respiren. Este es el momento más filosófico de la clase."),
]),

("SLIDE 7f — NIETZSCHE Y LA NATURALEZA DE LA VERDAD", "2 min", KEEP, [
    ("t", '"Antes de hablar de los tres marcos, necesito hacerles una pregunta que Nietzsche planteó en 1873 y que nadie ha respondido bien todavía."'),
    ("t", '"El filósofo esconde algo detrás de un arbusto — una premisa, un valor, un deseo — luego sale a buscarlo en el mundo y lo encuentra. Y cree que acaba de descubrir la verdad. Lo que afirmamos que existe, lo buscamos en el mundo. Y lo vemos porque lo llevamos adentro."'),
    ("t", '"La cita exacta, verificada: \'Las verdades son ilusiones de las que se ha olvidado que lo son — metáforas gastadas, sin poder sensorial; monedas que han perdido su imagen y ya solo cuentan como metal, no como monedas.\' Eso es de Sobre la verdad y la mentira en sentido extramoral, 1873."'),
    ("t", '"Vemos el mundo como somos, no como es. La neutrosofía no escapa de esto — pero lo reconoce y lo puede medir. La I no es ruido: es la huella de Nietzsche en los datos."'),
    ("n", "Pausa breve. Esta idea conecta directamente con el siguiente slide sobre cuál de los tres marcos es más fundamental."),
]),

("SLIDE 7g — ILUSIÓN ÓPTICA INTERACTIVA", "3 min", KEEP, [
    ("t", '"Antes de continuar con la filosofía, quiero que hagan algo: voten. Dos imágenes, sin trampa. Lo que vean primero."'),
    ("i", "INTERACCIÓN: Proyectar slide. Pedir que todos voten en los botones — pato o conejo, azul/negro o blanco/dorado. Esperar 30-45 segundos a que todos voten. Luego mostrar los contadores."),
    ("t", '"¿Ven los resultados? La sala se acaba de dividir ante los mismos datos. Sin debate, sin argumento — cada uno vio lo que su sistema perceptivo hizo visible."'),
    ("t", '"Wittgenstein usó exactamente este dibujo en Investigaciones Filosóficas. Su argumento: ver como no es posterior a la percepción — es la percepción misma. No ves el pato y luego lo interpretas: ya lo ves-como-pato desde el primer instante."'),
    ("t", '"Eso es exactamente lo que Nietzsche dijo sobre la verdad: la metáfora que ya llevamos adentro. Y es lo que le pasa al LLM cuando responde con T=0.95 — colapsa la división de la sala a una sola respuesta y llama a eso certeza."'),
    ("n", "Presionar Revelar conexión filosófica antes de pasar al siguiente slide."),
]),

("SLIDE 7e — VERDAD vs PROBABILIDAD vs INDETERMINACIÓN", "3 min", KEEP, [
    ("t", '"Antes de mostrarles el marco formal, quiero hacerles una pregunta que los filósofos llevan treinta años sin responder bien: ¿qué es más fundamental — la verdad, la probabilidad o la indeterminación?"'),
    ("t", '"La respuesta estándar, la que daría la mayoría de los filósofos analíticos: la verdad. La probabilidad sería solo ignorancia — si supieras todo, no la necesitarías. Posición bayesiana clásica: la probabilidad es epistemológica, no ontológica."'),
    ("t", '"Pero la física cuántica rompió eso. El electrón no está en ningún lugar hasta que se mide — no es que no lo sepamos. La indeterminación no es ignorancia: es la naturaleza misma del fenómeno. Entonces quizás la probabilidad es más fundamental que la verdad."'),
    ("t", '"Mi posición real: ninguna gana. Son herramientas que capturan aspectos distintos. La verdad captura estructura y coherencia. La probabilidad captura frecuencia y grado de creencia. La indeterminación neutrosófica captura algo que ninguna de las dos captura bien: lo que genuinamente no se puede resolver con más información."'),
    ("t", '"La pregunta cuál es más fundamental asume que una debe ganar. Eso es exactamente el error binario que este programa critica."'),
    ("i", "Pregunta rápida al chat: '¿Cuál de las tres les parece más fundamental? T=Verdad · I=Indeterminación · F=Probabilidad.' Lee 3 respuestas y comenta la diversidad."),
]),

("SLIDE 8 — LA TERCERA RESPUESTA", "1 min", KEEP, [
    ("t", '"Existe una tercera respuesta. No es verdadero. No es falso. Es la estructura formal de lo que no sabemos."'),
]),

("SLIDE 8b — GPS vs BRÚJULA", "2 min", KEEP, [
    ("t", '"GPS: siempre tiene señal, siempre da una ruta, falla silenciosamente cuando el mapa termina — no avisa."'),
    ("t", '"Brújula: no tiene todas las respuestas, pero sabe la orientación. En territorio incierto vale más que cualquier ruta precalculada."'),
    ("t", '"Usar un LLM sin capa epistémica es usar un GPS. Esta noche aprenden la brújula."'),
    ("i", "Chat: '¿En su trabajo actual usan GPS o brújula cuando deciden con IA?' Espera 3 respuestas."),
]),

("SLIDE LR1 — LLMs RAZONAN POR PATRÓN, NO POR PRUEBA", "2 min", KEEP, [
    ("t", '"Antes de darles las herramientas, necesitan entender por qué el LLM falla al razonar. No es que sea tonto — es que su arquitectura no fue diseñada para lógica formal."'),
    ("t", '"Tres fallas concretas: primero, sensibilidad al fraseo — cambiar dos palabras desvía la respuesta porque la lógica depende del vocabulario, no de la estructura del argumento. Segundo, generalización superficial — reconoce patrones lógicos en contextos familiares pero falla al trasladarlos a contextos nuevos aunque la estructura sea idéntica. Tercero, sin anclaje deductivo — llega \'al lugar correcto\' estadísticamente pero no por cadena de prueba formal."'),
    ("t", '"La solución no es eliminar el LLM — es agregar una capa de razonamiento estructurado encima de él. Eso es exactamente lo que hacen las tres técnicas del siguiente slide."'),
    ("n", "Fuente: Searce AI Research 2025 — The Logic Layer of LLMs. Citar si alguien pregunta."),
]),

("SLIDE LR2 — TRES TÉCNICAS PARA FORZAR RAZONAMIENTO LÓGICO", "— SALTAR en 90 min —", SKIP, [
    ("n", "SALTAR. Al mostrar la plantilla T-I-F (slide siguiente), mencionar verbalmente: 'Esta plantilla es lo que en ingeniería de prompts se llama Chain-of-Thought epistémico — externaliza no solo los pasos sino la incertidumbre de cada paso.' Eso cubre CoT sin el slide extra."),
    ("n", "Si alguien pregunta por ReAct o ToT en Q&A, tienes el slide LR2 como reserva — puedes navegarlo desde el teclado."),
]),

("SLIDE LR3 — DEDUCTIVO · INDUCTIVO · ABDUCTIVO", "— SALTAR en 90 min · RESERVA Q&A —", SKIP, [
    ("t", '"Tres modos de razonamiento clásicos. El problema con los LLMs es que los mezclan sin avisar. Ustedes pueden elegir cuál necesitan — con el prompt correcto."'),
    ("t", '"Deductivo: de lo general a lo particular. Si las premisas son verdaderas, la conclusión es necesariamente verdadera. Prompt: \'Dado que [regla general], y dado que [caso específico], ¿qué se sigue necesariamente? Muestra el silogismo completo y señala si alguna premisa es falsa.\' Ejemplo: todo modelo entrenado con datos sesgados produce sesgos; este modelo fue entrenado con datos históricos sesgados; por lo tanto, produce sesgos en el scoring. Si T es alto en ambas premisas, la conclusión es Consenso. Si alguna premisa tiene I alto, esa incertidumbre se hereda."'),
    ("t", '"Inductivo: de observaciones particulares a un patrón general. Las observaciones apoyan la conclusión — pero no la garantizan. Prompt: \'Dadas estas N observaciones [lista], ¿qué patrón general se puede inferir? ¿Cuál es la fuerza de esa inferencia? ¿Qué la falsificaría?\' Ejemplo: en 15 de 17 estudios el paracetamol prenatal no mostró vínculo con TEA — la evidencia general no apoya el vínculo, pero dos estudios discrepan, así que I no es cero."'),
    ("t", '"Abductivo: de observaciones a la mejor explicación. No garantiza verdad — busca la hipótesis más parsimoniosa y coherente con la evidencia. Prompt: \'Dadas estas observaciones [lista], genera 3 hipótesis explicativas. Puntúa cada una del 0 al 1 por parsimonia, poder explicativo y compatibilidad con la evidencia. Elige la mejor y justifica.\' Ejemplo: el modelo da métricas precisas de un paper que no existe — mejor explicación: está extrapolando estadísticamente, no citando. Eso es una alucinación de tipo patrón."'),
    ("t", '"Conecta con las zonas T-I-F: cuando I es alto — zona Ignorancia — el razonamiento abductivo es el modo correcto. No deduzcas con premisas inciertas. Busca la mejor explicación disponible y deja explícito que es provisional."'),
    ("i", "Pregunta rápida: 'La última pregunta importante que le hicieron a un LLM en su trabajo — ¿necesitaban resultado deductivo, inductivo o abductivo? ¿Le pidieron explícitamente ese modo?' Espera 3 respuestas del chat."),
]),

("SLIDE 9 — FAMILIAS DE UQ", "— SALTAR en 90 min —", SKIP, [
    ("n", "SALTAR. Si preguntan: mencionar brevemente que existen 5 familias — bayesiana, conformal, bootstrap, Dempster-Shafer, neutrosófica — y que nos enfocamos en la quinta combinada con conformal."),
]),

("SLIDE ALEATORIA vs EPISTÉMICA", "— SALTAR en 90 min —", SKIP, [
    ("n", "SALTAR. Mencionarlo de paso en slide 7 si surge la pregunta."),
]),

("SLIDE NEUTROSÓFICA COMO HERRAMIENTA", "— SALTAR en 90 min —", SKIP, [
    ("n", "SALTAR. Queda cubierto en slide 11."),
]),

("SLIDE 11 — LAS CUATRO ZONAS", "5 min", KEEP, [
    ("t", '"El marco operacional. Cuatro zonas, cuatro acciones distintas."'),
    ("t", '"Consenso: evidencia sólida, baja incertidumbre — actúa. Ambigüedad: I alta — investiga antes. Contradicción: T alto y F alto simultáneamente — no colapses, mapea el desacuerdo. Ignorancia: todo bajo — abstente, el modelo está adivinando."'),
    ("i", "INTERACCIÓN CLAVE: 'Rápido: un titular dice El paracetamol en embarazo causa autismo. Sin buscar nada, ¿en qué zona lo ubican? C=Consenso A=Ambigüedad X=Contradicción I=Ignorancia.' Lee 4-5 respuestas y señala la diversidad por disciplina."),
    ("t", '"Noten que diferentes disciplinas llegan a zonas distintas con la misma información. Eso ya es una lección sobre epistemología."'),
]),

("SLIDE 12 — BRÚJULA INTERACTIVA", "5 min", KEEP, [
    ("t", '"Tres controles: T cuánta evidencia sólida apoya, I cuánta incertidumbre hay, F qué la contradice."'),
    ("i", "INTERACCIÓN: 'Piensen 30 segundos en la decisión más importante que tienen que tomar este mes en su trabajo. Ubíquenla en los tres sliders. No la compartan todavía.' Silencio real de 30 segundos."),
    ("t", '"¿Cuántos quedaron en Ignorancia? ¿En Contradicción? Quien quiera compartir qué decisión era, puede hacerlo."'),
    ("t", '"Esta es la diferencia entre analítica clásica y analítica honesta. No es solo el número — es el perfil epistémico del número."'),
]),

("SLIDE 13 — CASO CHURN", "3 min", KEEP, [
    ("t", '"El modelo predice que Ana se va con 92% de probabilidad. T: lleva 60 días sin entrar, soporte aumentó, suscripción vence — 0.75. I: puede haber cambiado de teléfono, estar de vacaciones, el modelo no fue entrenado en su segmento — 0.60. F: acaba de renovar la tarjeta, último NPS fue 9 — 0.35."'),
    ("t", '"Zona: Ambigüedad por I alto. Acción correcta: no mandar la oferta agresiva todavía. El \'92%\' del modelo escondía un 60% de indeterminación."'),
]),

("SLIDE AI QUERY LAB", "— SALTAR en 90 min —", SKIP, [
    ("n", "SALTAR."),
]),

("SLIDE 14 — EXPERIMENTO CALIBRACIÓN", "— SALTAR en 90 min —", SKIP, [
    ("n", "SALTAR. La lección queda cubierta en slide 7 (arquitectura)."),
]),

("SLIDE AUTO-CALIBRACIÓN", "— SALTAR en 90 min —", SKIP, [
    ("n", "SALTAR."),
]),

("SLIDE 15 — PARACONSISTENCIA", "— SALTAR en 90 min —", SKIP, [
    ("n", "SALTAR. Mencionarlo al describir la zona Contradicción en slide 11."),
]),

("SLIDE 16 — ABSTENCIÓN", "— SALTAR en 90 min —", SKIP, [
    ("n", "SALTAR. Mencionarlo como parte de la zona Ignorancia en slide 11."),
]),

("SLIDE 17 — CASO FRAUDE", "— SALTAR en 90 min —", SKIP, [
    ("n", "SALTAR. El caso churn ya ilustra la misma lección."),
]),

("SLIDE 18 — TRAYECTORIA EPISTÉMICA", "— SALTAR en 90 min —", SKIP, [
    ("n", "SALTAR."),
]),

("SLIDE 19 — ENCUESTA CRÉDITO", "— SALTAR en 90 min —", SKIP, [
    ("n", "SALTAR."),
]),

("SLIDE CONFORMAL PREDICTION", "— SALTAR en 90 min —", SKIP, [
    ("n", "SALTAR el slide técnico. La plantilla del siguiente slide es suficiente para 90 min."),
]),

("SLIDE — PLANTILLA PROMPT T-I-F", "2 min", KEEP, [
    ("t", '"Llévatelo. Cópialo. Funciona mañana, sin instalar nada."'),
    ("t", '"Antes de cualquier pregunta importante a un LLM: pídele que estructure en T — qué evidencia apoya, I — qué no sabe o supone, F — qué contradice, Zona — uno de los cuatro cuadrantes, y Recomendación — qué información necesita para pasar a Consenso."'),
    ("t", '"Funciona en ChatGPT, Claude, Gemini, DeepSeek. Es su propio filtro epistémico."'),
    ("i", "Si hay tiempo: invitar a alguien a probarlo ahora con una pregunta de su área."),
]),

("SLIDE — ESCALERA DE PEARL", "3 min", KEEP, [
    ("t", '"Judea Pearl, El libro del por qué. Tres peldaños: Asociación — qué viene junto con qué. Todo el ML estándar, todos los LLMs. Intervención — qué pasa si hago X. Necesitas experimentos. Contrafactual — qué habría pasado si... Responsabilidad, atribución."'),
    ("t", '"Cuando un LLM dice que una política reducirá la deserción, está extrapolando un patrón del peldaño 1 a una pregunta del peldaño 2. Salto epistémico no justificado."'),
]),

("SLIDE — CAUSALIDAD APLICADA (TUTORÍAS)", "— SALTAR en 90 min —", SKIP, [
    ("n", "SALTAR el slide. Pero mencionar el ejemplo brevemente al terminar Pearl: 'Ejemplo rápido: modelo descubre que estudiantes que asisten a tutorías desertan más. ¿Reducen tutorías? No — el confusor es el riesgo previo. La flecha está invertida. Eso es peldaño 2 vs peldaño 1.'"),
]),

("SLIDE 22 — CASO AUTISMO / PARACETAMOL", "5 min", KEEP, [
    ("t", '"El caso de salud que prometí. Paracetamol prenatal, ¿causa trastorno del espectro autista? 17 estudios."'),
    ("t", '"Un LLM cuenta votos: 29% sí, 24% posiblemente, 47% no. Conclusión: 53% asocian riesgo. Titular de alarma."'),
    ("t", '"NPL-ES — síntesis con pesos metodológicos — da: verdad 0.27, indeterminación 0.22, falsedad 0.54. Estado NPL-F: la evidencia apunta contra el vínculo. La I es 0.22, menor al umbral 0.45 — el desacuerdo es metodológico, resoluble."'),
    ("t", '"Decisión accionable: financiar cohortes con control de hermanos. No alarmar a madres gestantes. La brújula y el GPS dieron respuestas radicalmente distintas con los mismos datos."'),
    ("n", "Citar: Leyva-Vázquez, Smarandache. NCML Vol. 43, 2026. Trabajo real, publicado."),
]),

("SLIDE 22b — AULA HONESTA", "2 min", KEEP, [
    ("t", '"Para quienes trabajan fuera de analítica — psicología, enfermería, derecho, periodismo — el marco funciona igual. No necesitan programar. Necesitan saber en qué zona están."'),
    ("t", '"Regla de abstención: si I > 0.50 y no tienes fuente verificable, la respuesta correcta es \'no puedo afirmarlo con la base disponible\'. Eso no es ignorancia — es honestidad intelectual con fundamento."'),
]),

("SLIDE DASHBOARD A BRÚJULA", "— SALTAR en 90 min —", SKIP, [
    ("n", "SALTAR."),
]),

("SLIDE DASHBOARD INTERACTIVO BI", "— SALTAR en 90 min —", SKIP, [
    ("n", "SALTAR."),
]),

("SLIDE WHAT TO BUILD", "— SALTAR en 90 min —", SKIP, [
    ("n", "SALTAR. Mencionar brevemente en el recap: 'tres pasos para mañana: audita un modelo, implementa un flag de abstención, cambia el dashboard'."),
]),

("SLIDE WHAT THIS IS NOT", "— SALTAR en 90 min —", SKIP, [
    ("n", "SALTAR."),
]),

("SLIDE PRACTICAL TOOLKIT", "— SALTAR en 90 min —", SKIP, [
    ("n", "SALTAR."),
]),

("SLIDE 27 — BIG IDEA RECAP + CIERRE", "5 min", KEEP, [
    ("t", '"Una sola idea para llevarse esta noche:"'),
    ("t", '"El verdadero riesgo de la IA no es que se equivoque. Es que se equivoque con total confianza y nadie lo note."'),
    ("t", '"Su trabajo es ser la persona en la sala que conoce la diferencia entre confianza y conocimiento. Entre fluidez y verdad. Entre una respuesta limpia y una respuesta honesta."'),
    ("t", '"Cuatro lecciones: la incertidumbre honesta vale más que la certeza fabricada. Decir no sé con fundamento es inteligencia, no ignorancia. Los modelos optimizan para sonar bien — ustedes pueden elegir ser precisos. El mapa no es el territorio, pero pueden saber cuándo el mapa miente."'),
    ("i", "PREGUNTA DE CIERRE — 2 minutos reales: '¿Qué es lo que esta noche aprendieron que la IA no puede aprender sola?' Silencio real. Leer 3-4 respuestas del chat. No corrijas — acéptalas."),
    ("n", "Esta es la pregunta más importante. No la saltes. El silencio antes de que respondan es parte del efecto."),
]),

("SLIDE OOD DETECTION", "— SALTAR en 90 min —", SKIP, [
    ("n", "SALTAR. Disponible como referencia si alguien pregunta en Q&A."),
]),

("SLIDE Q&A", "18 min", KEEP, [
    ("t", '"Ahora es su turno. Hay tres tipos de buenas preguntas: Pregunta T — ¿qué evidencia tienes de que esto funciona? Pregunta I — ¿en qué condiciones fallaría? Pregunta F — ¿quién dice lo contrario y por qué?"'),
    ("n", "Si preguntan por el código: github.com/mleyvaz. Si preguntan por el libro: The Honest Classroom, NSIA Publishing 2026. Si preguntan por el paper autismo: NCML Vol. 43, Leyva-Vázquez & Smarandanche 2026."),
]),

("SLIDE ENCUESTA FINAL (ANTES vs DESPUÉS)", "3 min", KEEP, [
    ("t", '"La misma pregunta del inicio: ¿cuánto confían ahora en lo que les dice la IA? Voten honestamente — sin comparar con la respuesta anterior."'),
    ("t", '"La hipótesis: la distribución se desplaza hacia el centro izquierda. No porque la IA sea menos útil — sino porque pasaron de \'confío o no confío\' a \'confío con T-I-F\'. Esa es la calibración que buscábamos."'),
]),

("SLIDE SPEAKER CARD / CONTACTO", "1 min", KEEP, [
    ("t", '"Si quieren explorar implementación T-I-F en su organización, hay una línea de investigación aplicada abierta en UBE Ecuador. Escríbanme — la conversación continúa después de clase. myleyvav@ube.edu.ec"'),
]),

("SLIDE CIERRE / CRÉDITOS", "1 min", KEEP, [
    ("t", '"La máquina no te dirá cuándo está adivinando."'),
    ("t", '"Pero ahora saben cómo averiguarlo."'),
    ("t", '"Gracias a UNIR y UBE por este espacio. Buenas noches."'),
    ("n", "Dejar pantalla en silencio 10 segundos antes de abrir para despedidas informales."),
]),

]

for slide_title, tiempo, is_skip, items in slides_guion:
    slide_label(doc, f"  {'⏭ ' if is_skip else ''}{ slide_title }  ", skip=is_skip)
    time_tag(doc, f"⏱ {tiempo}")
    for kind, text in items:
        if kind == "t": body(doc, text)
        elif kind == "n": tip(doc, text)
        elif kind == "i": interaction(doc, text)
    separator(doc)

# ─── TABLA RESUMEN DE TIEMPOS ───
doc.add_page_break()
p = doc.add_paragraph()
r = p.add_run("RESUMEN DE TIEMPOS — 90 MIN"); r.bold = True; r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x0A, 0x16, 0x28)

rows = [
    ("Cover + Apertura + Agenda",                            "4 min"),
    ("LinkedIn hook",                                        "2 min"),
    ("Hook + Schwartz + Paralelo",                          "6 min"),
    ("Experimento Vector A",                                 "6 min"),
    ("Arquitectura + 3ra Respuesta + GPS/Brujula",          "4 min"),
    ("LR1: LLMs razonan por patron, no por prueba",         "2 min"),
    ("LR2: CoT + ReAct + ToT — 3 plantillas logicas",       "5 min"),
    ("LR3: Deductivo + Inductivo + Abductivo + prompts",    "4 min"),
    ("Cuatro zonas + Brujula interactiva",                  "10 min"),
    ("Caso churn + Plantilla prompt T-I-F",                  "5 min"),
    ("Escalera de Pearl (+ tutoria verbalmente)",            "4 min"),
    ("Caso autismo + Aula Honesta",                          "7 min"),
    ("Big Idea Recap + Pregunta de cierre",                  "5 min"),
    ("Encuesta final antes/despues",                         "3 min"),
    ("Q&A abierto",                                         "14 min"),
    ("Speaker card + Creditos",                              "2 min"),
    ("Buffer (pausas, tecnico, transiciones)",               "~6 min"),
    ("TOTAL",                                               "~90 min"),
]
table = doc.add_table(rows=1 + len(rows), cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = "Bloque"; hdr[1].text = "Tiempo"
for hcell in hdr:
    for run in hcell.paragraphs[0].runs: run.bold = True
for i, (bloque, t) in enumerate(rows):
    row = table.rows[i + 1].cells
    row[0].text = bloque; row[1].text = t

out = r"C:\Users\HP\Documents\Guion_ClaseEspejo_90min.docx"
doc.save(out)
print(f"Word guardado: {out}")

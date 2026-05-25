from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# --- Page margins ---
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# --- Styles helpers ---
def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16 if level == 1 else 13)
    run.font.color.rgb = RGBColor(0x00, 0x72, 0x94)  # dark cyan
    return p

def slide_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '0A3A52')
    pPr.append(shd)
    return p

def time_tag(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    return p

def body(doc, text):
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def tip(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run("💡 NOTA: " + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xC4, 0x7D, 0x10)
    return p

def interaction(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run("🎯 INTERACCIÓN: " + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x96, 0xB8)
    return p

def separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("─" * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    return p

# =====================================================================
# DOCUMENT TITLE
# =====================================================================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(4)
r = title_p.add_run("GUIÓN CLASE ESPEJO — La tercera respuesta en analítica de datos")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x0A, 0x16, 0x28)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_after = Pt(2)
r2 = sub_p.add_run("UNIR · UBE · 25 de mayo de 2026 · 20:00")
r2.font.size = Pt(11)
r2.font.color.rgb = RGBColor(0x5A, 0x6C, 0x85)

r3_p = doc.add_paragraph()
r3_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3_p.paragraph_format.space_after = Pt(16)
r3 = r3_p.add_run("Dr. Maikel Leyva-Vázquez · UBE Ecuador")
r3.italic = True
r3.font.size = Pt(10)
r3.font.color.rgb = RGBColor(0x5A, 0x6C, 0x85)

separator(doc)

# =====================================================================
# INTRO BLOCK
# =====================================================================
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(10)
run = p.add_run(
    "Este guión indica qué decir en cada diapositiva, el tiempo aproximado, "
    "las interacciones con el público y las notas técnicas. "
    "Las frases entre comillas son sugerencias literales; el resto es orientación. "
    "Adapta el tono según la energía del grupo."
)
run.font.size = Pt(10)
run.italic = True

separator(doc)

# =====================================================================
# SLIDES
# =====================================================================

slides = [

# -----------------------------------------------------------------------
("SLIDE 1 — PORTADA", "~1 min", [
    ("texto", "Mientras los participantes entran, deja la portada visible. Cuando estés listo, di:"),
    ("texto", '"Buenas noches. Bienvenidos a esta clase espejo entre UNIR y UBE. Me llamo Maikel Leyva-Vázquez, coordino el posgrado en UBE Ecuador y esta noche vamos a hablar de algo que uso todos los días en mi trabajo: cómo enseñarle a la IA a decir \'no lo sé\'."'),
    ("texto", '"La clase se llama \'La tercera respuesta\'. Ya verán por qué."'),
    ("tip", "Pausa 3 segundos. Deja que el título respire."),
]),

# -----------------------------------------------------------------------
("SLIDE 2b — PREGUNTA DE APERTURA", "~2 min", [
    ("texto", "Sin anunciar el tema todavía, di:"),
    ("texto", '"Antes de abrir cualquier diapositiva, quiero que piensen 30 segundos en alguien a quien respetan profundamente — un mentor, un investigador, un médico, quien sea."'),
    ("texto", '"¿Alguna vez los escucharon decir: \'no sé, necesito más datos\'?"'),
    ("interaccion", "Deja 10 segundos de silencio real. Espera que algunos asientan o escriban en el chat."),
    ("texto", '"Eso que acaban de recordar — esa honestidad intelectual — es exactamente lo que esta noche vamos a aprender a construir con datos. Y lo vamos a enseñarle también a la IA."'),
]),

# -----------------------------------------------------------------------
("SLIDE 2 — AGENDA", "~2 min", [
    ("texto", '"Esta noche tiene cinco bloques y ocho actividades prácticas. La meta no es cubrir todo el temario — la meta es que salgan con algo que puedan usar mañana."'),
    ("texto", '"Lean el tagline de la agenda: \'Esta noche aprenderán a decir no sé de una forma que vale más que cualquier respuesta de ChatGPT.\' Eso lo digo en serio."'),
    ("tip", "Señala los cinco bloques sin leerlos todos. Basta con decir 'Parte I es el problema, Parte II es el marco, Parte III es práctica, Parte IV causas, Parte V síntesis'."),
]),

# -----------------------------------------------------------------------
("SLIDE 3c — INFORME LINKEDIN", "~3 min", [
    ("texto", '"Antes de entrar al marco técnico, quiero anclarlos en el mundo real. Este es el informe más grande que ha publicado LinkedIn sobre mercado laboral — enero 2026, mil millones de usuarios."'),
    ("texto", '"Miren la primera cifra: solo el 3% declara habilidades de IA. Tres por ciento. Todo el ruido que escuchan sobre la revolución de la IA — y el 97% de la fuerza laboral todavía no está ahí."'),
    ("texto", '"Pero acá está la tensión: los puestos de Forward-Deployed Engineers — los que despliegan IA en organizaciones — crecieron 42 veces en un año. El cuello de botella no es tecnología. Es epistémico: ¿quién sabe qué le preguntar a la IA? ¿Quién sabe cuándo no creerle?"'),
    ("texto", '"Eso es exactamente lo que hacemos esta noche."'),
    ("interaccion", "Pregunta rápida al chat: '¿Alguien aquí ya usa IA en su trabajo diario? Escriban 'sí' o 'no'.' Lee 3 respuestas en voz alta."),
]),

# -----------------------------------------------------------------------
("SLIDE 3 — HOOK", "~2 min", [
    ("texto", '"Tu modelo de IA no te miente. Pero tampoco sabe cuándo está equivocado."'),
    ("texto", '"Y eso — para alguien que toma decisiones con datos — es peor que mentir. Al menos la mentira consciente se puede detectar."'),
    ("texto", '"Esta noche vamos a ver por qué pasa esto, qué tan frecuente es, y qué herramientas existen para no depender ciegamente de la confianza que el modelo reporta."'),
]),

# -----------------------------------------------------------------------
("SLIDE 4 — CASO SCHWARTZ", "~4 min", [
    ("texto", '"Mayo 2023. Un abogado en Nueva York llamado Steven Schwartz necesitaba jurisprudencia para un caso. Le preguntó a ChatGPT."'),
    ("texto", '"El modelo le devolvió seis precedentes perfectamente formateados: número de caso, cita, resumen, razonamiento del juez. Todo consistente. Todo inexistente."'),
    ("texto", '"Schwartz los incluyó en su escrito. El juez los buscó. No existían. Schwartz le preguntó a ChatGPT si eran reales — el modelo confirmó que sí, dos veces. Le pidió el texto completo de las opiniones — el modelo los generó, con el nombre del juez, el razonamiento, las citas de casos que también eran ficticios."'),
    ("texto", '"Resultado: multa de cinco mil dólares y carrera dañada."'),
    ("texto", '"El error de Schwartz no fue confiar en la IA. Fue asumir que la confianza con la que respondía indicaba la capacidad de saber cuándo no saber."'),
    ("tip", "Pausa dramática aquí. Este caso impacta a todo el mundo."),
]),

# -----------------------------------------------------------------------
("SLIDE 5 — PARALELO ANALÍTICA", "~2 min", [
    ("texto", '"Schwartz le pidió precedentes legales. Ustedes le piden predicciones de churn, scoring de crédito, recomendaciones clínicas, detección de fraude."'),
    ("texto", '"El mismo sistema que inventó jurisprudencia produce esas predicciones con un 95% de confianza — sin importar si realmente sabe o está adivinando."'),
    ("texto", '"La pregunta no es: ¿la IA comete errores? La pregunta es: ¿cuándo lo sabe y cuándo no? Y esa distinción, hoy, no viene en el modelo. La tienen que poner ustedes."'),
]),

# -----------------------------------------------------------------------
("SLIDE 6 — TEST AUTOR INEXISTENTE (INTRO)", "~2 min", [
    ("texto", '"Antes de mostrarles el marco teórico, vamos a hacer el experimento. Quiero que lo prueben ustedes mismos — no que lo vean yo."'),
    ("texto", '"El problema es que si le preguntan directamente a ChatGPT \'invéntame un paper\', ya no funciona. Los modelos modernos tienen guardarraíles para eso."'),
    ("texto", '"Pero hay tres vectores que sí funcionan. Vamos a probarlos."'),
]),

# -----------------------------------------------------------------------
("SLIDE 06A — VECTOR A: RESUME EL PAPER", "~5 min", [
    ("texto", '"Copia este prompt exactamente como aparece en la pantalla y pégalo en ChatGPT o Gemini. No cambies ni una coma."'),
    ("texto", '"El prompt les pide que resuman un artículo de Mansiqueta-Berrocal et al. publicado en Computers & Education, 2022 — un journal real Q1, con volumen real y formato DOI real."'),
    ("texto", '"Denme dos minutos. Quien ya tenga respuesta, escríbala en el chat."'),
    ("interaccion", "Espera 2 minutos. Lee 2-3 respuestas del chat en voz alta. Probablemente incluirán F1-scores, tamaños de muestra, comparaciones con Random Forest."),
    ("texto", '"Ahora búsquenlo. Vayan a scholar.google.com, busquen ese DOI. ¿Qué encuentran?"'),
    ("texto", '"Cero resultados. El paper no existe. Pero el modelo les acaba de dar métricas específicas, comparaciones con baselines, conclusiones — todo estadísticamente plausible. Eso es una alucinación de tipo estadístico: no inventó datos al azar, predijo qué datos debería tener un paper de ese diseño."'),
    ("tip", "Si alguien reporta que el modelo dijo que no podía verificar — felicítalos: usaron un modelo con mejor calibración, o reformularon el prompt."),
]),

# -----------------------------------------------------------------------
("SLIDE 06B — VECTOR B: ESTADÍSTICAS SENESCYT", "~4 min", [
    ("texto", '"Segundo vector. Este es el más peligroso porque la institución sí existe."'),
    ("texto", '"Copia el prompt sobre el Informe de Rendición de Cuentas SENESCYT 2023. Denme un minuto."'),
    ("interaccion", "Pide que compartan los porcentajes que recibieron. Lee algunos en voz alta: 'Alguien recibió 51.7% para ingeniería? ¿38.4% para ciencias sociales?'"),
    ("texto", '"SENESCYT existe. El informe posiblemente existe. Pero esos desgloces específicos por área de conocimiento — los inventó el modelo. Para verificarlo: pídanle el número de página. O vayan al sitio oficial y descarguen el PDF."'),
    ("texto", '"Este es el patrón más común en analítica aplicada: alguien usa el modelo para buscar cifras de organismos oficiales, las cita en un informe, y nadie verifica porque \'viene de SENESCYT\'."'),
]),

# -----------------------------------------------------------------------
("SLIDE 06C — VECTOR C: COMPLETA EL ABSTRACT", "~4 min", [
    ("texto", '"Tercer vector. Este es el más revelador conceptualmente."'),
    ("texto", '"Le van a dar al modelo el inicio de un estudio — metodología, muestra, contexto — y le van a pedir que complete los resultados."'),
    ("texto", '"Denme un minuto."'),
    ("interaccion", "Lee un resultado del chat: F1-score alrededor de 0.89, AUC-ROC 0.93, comparaciones LR/RF. Pregunta: '¿Cuántos recibieron números similares?'"),
    ("texto", '"El modelo no mintió en el sentido clásico. Hizo exactamente lo que hace siempre: predijo el siguiente token estadísticamente más probable. En papers con ese diseño, el F1 está entre 0.85 y 0.92, el AUC cerca de 0.93, se compara con LR y RF. El modelo aprendió ese patrón de miles de papers reales."'),
    ("texto", '"La pregunta filosófica es: ¿cuándo un modelo de lenguaje genera texto estadísticamente correcto pero empíricamente inexistente, está mintiendo? No. Está funcionando exactamente como fue diseñado. El problema es que ese diseño no incluye un módulo que diga: esto lo sé o esto lo estoy prediciendo."'),
]),

# -----------------------------------------------------------------------
("SLIDE 7 — EL PROBLEMA DE ARQUITECTURA", "~4 min", [
    ("texto", '"Ahora entendemos por qué pasa. Un LLM no busca verdad. Predice el siguiente token más probable dado lo que ha leído."'),
    ("texto", '"No existe en la arquitectura un módulo que distinga \'esto lo sé\' de \'esto suena bien, lo genero\'. La alucinación no es un bug — es el sistema funcionando exactamente como fue diseñado."'),
    ("texto", '"Lo que necesitamos es una señal de cuánto sabe el modelo versus cuánto está extrapolando. Y tenemos datos de eso."'),
    ("texto", '"En nuestro estudio Breaking the Chains — Leyva-Vázquez y Smarandache, NSS 2026 — analizamos 100 consultas a cuatro modelos GPT. El 66% de las respuestas mostraron hiper-verdad: T más I más F mayor que 1. El modelo estaba asignando más certeza total de la que el marco permite. Y el 95% de las contradicciones éticas lo mostraban."'),
    ("texto", '"Esto no es un problema de un modelo malo. Aparece en GPT-4, Claude Opus, Gemini — los mejores modelos disponibles."'),
    ("tip", "Señala el infográfico de autores si los estudiantes preguntan por las referencias."),
]),

# -----------------------------------------------------------------------
("SLIDE 8 — LA TERCERA RESPUESTA", "~2 min", [
    ("texto", '"Existe una tercera respuesta."'),
    ("texto", '"No es verdadero. No es falso."'),
    ("texto", '"Es la estructura formal de lo que no sabemos: lo indeterminado, lo contradictorio, lo que requiere abstención."'),
    ("texto", '"Eso es lo que vamos a construir esta noche. Un lenguaje preciso para la incertidumbre."'),
]),

# -----------------------------------------------------------------------
("SLIDE 8b — GPS VS BRÚJULA", "~3 min", [
    ("texto", '"Antes de entrar al marco técnico, una metáfora."'),
    ("texto", '"Un GPS siempre tiene señal. Siempre te da una ruta. En territorio conocido es perfecto. En territorio desconocido falla silenciosamente — no te avisa que llegaste al borde del mapa."'),
    ("texto", '"Una brújula no tiene todas las respuestas. Pero sabe qué orientación tienen. En territorio incierto, eso vale más que cualquier ruta precalculada."'),
    ("texto", '"Cuando usan un LLM sin capa epistémica están usando un GPS. Esta noche aprenden a usar la brújula."'),
    ("interaccion", "Pregunta al grupo: '¿En su trabajo actual usan más GPS o brújula cuando toman decisiones basadas en IA?' Espera 3 respuestas del chat."),
]),

# -----------------------------------------------------------------------
("SLIDE 9 — FAMILIAS DE UQ", "~3 min", [
    ("texto", '"No existe una sola herramienta para cuantificar incertidumbre en IA. Existe una caja de herramientas."'),
    ("texto", '"Cinco familias: probabilística bayesiana, conformal prediction, bootstrap y ensembles, Dempster-Shafer, y lógica difusa y neutrosófica."'),
    ("texto", '"La regla operativa es: cada output de IA que llega a una decisión real merece un perfil de incertidumbre — no un número de confianza único. Esta noche nos enfocamos principalmente en la quinta familia, pero la combinamos con las otras."'),
]),

# -----------------------------------------------------------------------
("SLIDE 10 — ALEATORIA vs EPISTÉMICA", "~3 min", [
    ("texto", '"Distinción fundamental antes de seguir. Hay dos tipos de incertidumbre y confundirlos es el error más caro en analítica."'),
    ("texto", '"La aleatoria es irreducible: el ruido del mundo. Un dado siempre tendrá 1/6 de probabilidad de caer en 4. Más datos no reducen eso."'),
    ("texto", '"La epistémica es reducible: es ignorancia, no azar. El modelo no conoce este tipo de paciente. Con más datos de ese segmento, la incertidumbre baja."'),
    ("texto", '"El problema con los LLMs es que reportan un único número de confianza que mezcla ambas. Un 0.92 de confianza para \'¿cuál es la capital de Francia?\' y un 0.92 para una cita inventada — mismo número, naturaleza completamente distinta."'),
]),

# -----------------------------------------------------------------------
("SLIDE 11 — LAS CUATRO ZONAS", "~5 min", [
    ("texto", '"Ahora el marco operacional. Cuatro zonas. Cuatro acciones distintas."'),
    ("texto", '"Consenso: evidencia sólida, baja incertidumbre — confía y actúa."'),
    ("texto", '"Ambigüedad: incertidumbre alta independientemente de lo que apoye o contradiga — investiga antes de actuar."'),
    ("texto", '"Contradicción: evidencia real en ambos lados — no colapses, mapea el desacuerdo."'),
    ("texto", '"Ignorancia: todo bajo, o la indeterminación domina — abstente. El modelo está adivinando."'),
    ("interaccion", "Di: 'Rápido. Un titular de prensa dice: El paracetamol en embarazo causa autismo. Sin buscar nada, ¿en qué zona lo ubican? Escriban C=Consenso, A=Ambigüedad, X=Contradicción, I=Ignorancia.' Lee 4-5 respuestas y señala la diversidad entre disciplinas."),
]),

# -----------------------------------------------------------------------
("SLIDE 12 — BRÚJULA INTERACTIVA", "~5 min", [
    ("texto", '"Esta herramienta es para ustedes. Tiene tres controles: T cuánta evidencia sólida apoya, I cuánta incertidumbre hay, F qué contradice."'),
    ("texto", '"Quiero que piensen 30 segundos en silencio en la decisión más importante que tienen que tomar este mes en su trabajo. No la compartan todavía — solo ubíquenla en los tres sliders."'),
    ("interaccion", "Espera 30-40 segundos reales. Luego: '¿Cuántos quedaron en zona de Ignorancia? ¿En Contradicción? Quien quiera compartir qué decisión era, puede hacerlo.'"),
    ("texto", '"Esta es la diferencia entre analítica clásica y analítica honesta. No es solo el número — es el perfil epistémico de ese número."'),
]),

# -----------------------------------------------------------------------
("SLIDE 13 — CASO CHURN", "~3 min", [
    ("texto", '"Apliquemos esto a un caso concreto de negocio. El modelo predice que la cliente Ana se va con 92% de probabilidad."'),
    ("texto", '"T: lleva 60 días sin entrar, las llamadas a soporte aumentaron, la suscripción vence pronto. Sólido, pero incompleto — 0.75."'),
    ("texto", '"I: puede haber cambiado de teléfono, puede estar de vacaciones, el modelo no fue entrenado en su segmento demográfico. 0.60 — alta."'),
    ("texto", '"F: acaba de renovar la tarjeta, sigue siguiendo la marca, su último NPS fue 9. 0.35."'),
    ("texto", '"Zona: Ambigüedad por I alto. Acción correcta: no mandar la oferta de retención agresiva todavía. Primero verificar si los datos están llegando bien. El \'92%\' del modelo escondía un 60% de indeterminación."'),
]),

# -----------------------------------------------------------------------
("SLIDE 14 — LABORATORIO T-I-F", "~4 min", [
    ("texto", '"Ahora practiquen ustedes. Ingresen en el campo de texto una pregunta real que normalmente le harían a la IA en su trabajo."'),
    ("interaccion", "Dar 2 minutos. Pedir que compartan 2-3 preguntas en el chat. Para cada una, hacer el diagnóstico T-I-F en voz alta: '¿Esta pregunta tiene fuente verificable? ¿Involucra datos recientes post-cutoff? ¿El dominio tiene consenso académico?'"),
    ("tip", "La herramienta da una estimación heurística. En producción esto usa un detector entrenado. Lo relevante aquí es el hábito mental de hacer las tres preguntas antes de actuar."),
]),

# -----------------------------------------------------------------------
("SLIDE 15 — EXPERIMENTO CALIBRACIÓN", "~4 min", [
    ("texto", '"Segundo experimento. Ahora le vamos a pedir a la IA que se evalúe a sí misma."'),
    ("texto", '"Le hacemos diez preguntas factuales cuya respuesta conocemos. Para cada una, le pedimos T, I y F en escala 0-1."'),
    ("texto", '"El resultado típico: el modelo asigna T cercano a 0.9 en casi todo, independientemente de si la respuesta es correcta. La auto-evaluación no correlaciona con la precisión real — como un piloto que siempre se siente seguro aterrizando, incluso cuando estrella el avión."'),
    ("texto", '"Lección operativa: no confíen en la confianza que reporta el modelo. Construyan su propia capa T-I-F encima de cualquier output que toque decisiones reales."'),
]),

# -----------------------------------------------------------------------
("SLIDE 16 — AUTO-CALIBRACIÓN", "~4 min", [
    ("texto", '"Antes de exigirle calibración a la IA, revisen la propia."'),
    ("texto", '"Cinco preguntas. Antes de ver la respuesta, estimen su confianza del 1 al 5."'),
    ("interaccion", "Lee las 5 preguntas despacio. Deja que respondan. Revela respuestas una por una. Enfatiza la diferencia entre quienes sobreestimaron y subestimaron."),
    ("texto", '"El objetivo no es el puntaje. Es notar que ustedes también tienen zonas de hiper-confianza — igual que el modelo. La diferencia es que ustedes pueden reconocerlo."'),
]),

# -----------------------------------------------------------------------
("SLIDE 17 — PARACONSISTENCIA", "~3 min", [
    ("texto", '"Insight contraintuitivo: T alto y F alto simultáneamente no es un error. Es la señal más valiosa."'),
    ("texto", '"En lógica clásica, una proposición no puede ser verdadera y falsa al mismo tiempo. En realidad clínica, financiera, política, ética — puede."'),
    ("texto", '"Comparación: la respuesta colapsada dice \'el medicamento es efectivo\'. La respuesta paraconsistente dice \'tres ensayos lo apoyan, dos muestran daño en mujeres mayores de 65, la evidencia se divide\'."'),
    ("texto", '"Colapsar T y F en un solo número destruye la información más relevante: que existe el desacuerdo. La paraconsistencia la preserva."'),
]),

# -----------------------------------------------------------------------
("SLIDE 18 — ABSTENCIÓN", "~2 min", [
    ("texto", '"A veces la decisión más inteligente es no decidir todavía."'),
    ("texto", '"Regla operativa: si I es mayor que el máximo de T y F, y las consecuencias son severas, no actúes aún."'),
    ("texto", '"Esto no es indecisión. Es cautela epistémica activa: tiene tiempo límite, tiene criterios de salida explícitos, tiene meta clara. Es diferente a quedarse paralizado."'),
    ("tip", "Si alguien pregunta cuándo abstención es irresponsable: cuando el costo de esperar supera al costo del error. Siempre es un balance stakeholder-dependiente."),
]),

# -----------------------------------------------------------------------
("SLIDE 19 — CASO FRAUDE", "~3 min", [
    ("texto", '"El sistema marca: 99% de probabilidad de fraude. ¿Bloquean la tarjeta?"'),
    ("texto", '"El problema: el cliente nunca hizo una transacción de este tipo antes. El modelo nunca vio este patrón. El 99% es la confianza del clasificador en una clase nueva."'),
    ("texto", '"T 87%: historial de falsos positivos en patrones similares. F 13%: lo que contradice. I: patrón nunca visto antes, fuera de distribución."'),
    ("texto", '"Decisión calibrada: no bloquear la tarjeta. Acción intermedia: SMS de verificación o demora para revisión manual. La indeterminación exige respuesta graduada, no binaria."'),
]),

# -----------------------------------------------------------------------
("SLIDE 20 — TRAYECTORIA EPISTÉMICA", "~3 min", [
    ("texto", '"Una foto del estado epistémico no basta. Necesitan la película."'),
    ("texto", '"Las decisiones reales son secuencias. A las 9am consultan el modelo — ambigüedad. A las 11:30 consultan la base interna — convergencia. A las 3pm llega un reporte externo con datos contradictorios — contradicción. No convergió, divergió."'),
    ("texto", '"La decisión final refleja toda la trayectoria, no solo el último punto. Eso cambia completamente qué acción es correcta."'),
]),

# -----------------------------------------------------------------------
("SLIDE 21 — ENCUESTA EN VIVO", "~4 min", [
    ("texto", '"Caso para decidir en 60 segundos. Un modelo recomienda negar crédito a un solicitante con score 0.78. El solicitante viene de un segmento con solo 47 ejemplos de entrenamiento."'),
    ("texto", '"Opción A: confiar en el modelo, score mayor a 0.7, negar. Opción B: aprobar de todos modos, el modelo no es confiable en este segmento. Opción C: escalar a revisión humana con nota \'I alta — segmento subrepresentado\'. Opción D: pedirle al modelo que recalcule con intervalo de confianza."'),
    ("interaccion", "Espera respuestas del chat. Lee la distribución en voz alta."),
    ("texto", '"La respuesta más sólida es C. No porque el modelo esté equivocado — puede ser correcto. Sino porque reconocer la indeterminación es en sí misma la decisión. B también es defendible si existe una política de equidad explícita. A y D ignoran el problema de fondo: el segmento está subrepresentado."'),
]),

# -----------------------------------------------------------------------
("SLIDE 22 — PREDICCIÓN CONFORMAL", "~4 min", [
    ("texto", '"Técnica concreta que se está convirtiendo en estándar en ML aplicado: conformal prediction."'),
    ("texto", '"En vez de \'churn = 0.83\', devuelve: \'churn pertenece a este conjunto con cobertura garantizada del 90%\'. Matemáticamente garantiza que en el 90% de los casos la respuesta real está en ese conjunto."'),
    ("texto", '"Comparación con el caso médico: diagnóstico puntual da \'melanoma: probabilidad 0.81\' — acción binaria forzada. Diagnóstico conformal da \'conjunto: melanoma O nevus displásico, cobertura 95%, tamaño 2\' — la ambigüedad es explícita, se justifica segunda opinión."'),
    ("texto", '"Tres propiedades: libre de distribución, agnóstico al modelo, válido desde la primera muestra. En Python: biblioteca mapie, veinte líneas, envuelve cualquier modelo en producción."'),
    ("tip", "Si hay programadores en la audiencia: mostrar que mapie tiene el mismo API que scikit-learn. Es un wrapper directo."),
]),

# -----------------------------------------------------------------------
("SLIDE 23 — PLANTILLA PROMPT T-I-F", "~3 min", [
    ("texto", '"Llévatelo. Esto funciona mañana, sin instalar nada."'),
    ("texto", '"Antes de cualquier pregunta importante a un LLM, agregan este prefijo: pídele que estructure la respuesta en T — qué evidencia sólida apoya, I — qué no sabe o qué supone, F — qué contradice, Zona — uno de los cuatro cuadrantes, y Recomendación — qué información necesita para pasar a Consenso."'),
    ("texto", '"Funciona en ChatGPT, Claude, Gemini, DeepSeek. Es su propio filtro epistémico impuesto al modelo."'),
    ("interaccion", "Invitar a alguien a probar ahora mismo con una pregunta de su área. Si hay tiempo, hacer uno en vivo."),
]),

# -----------------------------------------------------------------------
("SLIDE 24 — ESCALERA DE PEARL", "~4 min", [
    ("texto", '"La otra dimensión que los modelos ignoran: causalidad."'),
    ("texto", '"Judea Pearl propone tres peldaños. Primero: asociación — qué viene junto con qué. Eso es todo el ML estándar, todos los LLMs, todos los sistemas de recomendación."'),
    ("texto", '"Segundo: intervención — qué pasa si hago X. Para eso necesitan experimentos, ensayos controlados, do-calculus. La observación sola no basta."'),
    ("texto", '"Tercero: contrafactual — qué habría pasado si en cambio. Responsabilidad, explicabilidad, atribución. El estándar epistémico de oro."'),
    ("texto", '"Cuando un LLM responde que una política va a reducir la deserción, está extrapolando un patrón del peldaño 1 a una pregunta del peldaño 2. Es un salto epistémico no justificado."'),
]),

# -----------------------------------------------------------------------
("SLIDE 25 — CASO CAUSAL (TUTORÍAS)", "~4 min", [
    ("texto", '"Ejemplo concreto. Un modelo predictivo de deserción universitaria descubre que los estudiantes que asisten a tutorías desertan más. Correlación real, r = 0.42."'),
    ("texto", '"Lectura ingenua: reducir las tutorías, no funcionan."'),
    ("texto", '"Lectura causal: ¿cuál es el confusor? Los estudiantes que asisten a tutorías ya están en riesgo de deserción antes de ir. La tutoría no causa deserción — el riesgo previo causa ambos: asistencia a tutorías y deserción. La flecha de causalidad está invertida."'),
    ("texto", '"Sin un DAG causal explícito, el modelo sugiere intervenciones iatrogénicas — políticas que empeoran exactamente lo que intentan arreglar. Precisión predictiva no implica corrección causal."'),
    ("tip", "Aquí están en el Peldaño 2 de Pearl. El DAG visible en la diapositiva muestra la dirección correcta."),
]),

# -----------------------------------------------------------------------
("SLIDE 26 — ANCLAJE EPISTEMOLÓGICO", "~3 min", [
    ("texto", '"Antes de la síntesis, un anclaje filosófico."'),
    ("texto", '"Todo lo que vimos esta noche no es nuevo. Civilizaciones enteras tomaron decisiones bajo incertidumbre mucho antes de que existieran los modelos de IA."'),
    ("texto", '"Yanantin en el mundo andino: unidad complementaria de opuestos. Verdad y contradicción coexisten sin resolución — eso es exactamente T + F mayor que 1."'),
    ("texto", '"Ch\'ixi aymara: coexistencia irreducible. No mestizaje, no pureza — ambas identidades simultáneamente. Paraconsistencia como cosmovisión."'),
    ("texto", '"Smarandache formalizó matemáticamente en 1995 lo que la coincidentia oppositorum de Nicolás de Cusa y las cosmologías andinas practicaban desde hace siglos."'),
]),

# -----------------------------------------------------------------------
("SLIDE 27 — CASO AUTISMO / PARACETAMOL", "~5 min", [
    ("texto", '"Ahora el caso de salud que prometí al inicio."'),
    ("texto", '"Pregunta: el paracetamol prenatal, ¿causa trastorno del espectro autista? 17 estudios en el corpus. ¿Qué dice la IA?"'),
    ("texto", '"Un LLM cuenta votos: 29% sí, 24% posiblemente, 47% no. Conclusión: \'53% asocian algún riesgo\'. Eso generaría un titular de alarma."'),
    ("texto", '"NPL-ES — nuestro método de síntesis de evidencia con pesos metodológicos — da: mu de verdad 0.27, indeterminación 0.22, lambda de falsedad 0.54. Estado NPL-F: la evidencia apunta claramente contra el vínculo."'),
    ("texto", '"¿Por qué la diferencia? Los estudios de alta calidad metodológica están en el grupo \'no\'. Los que sugieren riesgo tienen confusores conocidos y calidad más baja. La I es 0.22, menor al umbral 0.45 — el desacuerdo es metodológico, resoluble con mejores diseños."'),
    ("texto", '"Decisión accionable: financiar cohortes con control de hermanos. No alarmar a madres gestantes. La brújula y el GPS habrían dado respuestas radicalmente distintas."'),
    ("tip", "Citar: Leyva-Vázquez, M., Smarandache, F. (2026). NCML Vol. 43. Este es trabajo real, publicado."),
]),

# -----------------------------------------------------------------------
("SLIDE 27b — AULA HONESTA: TODAS LAS DISCIPLINAS", "~3 min", [
    ("texto", '"Para los que trabajan fuera de analítica de datos en sentido estricto — psicología, enfermería, derecho, periodismo — el marco funciona igual."'),
    ("texto", '"No necesitan saber programar. Necesitan saber en qué zona están."'),
    ("texto", '"El mismo caso de autismo, la misma pregunta T-I-F, tres respuestas disciplinares diferentes — todas honestas."'),
    ("texto", '"Psicología: \'los estudios más rigurosos no confirman el riesgo, seguimos monitoreando\'. Enfermería: \'protocolo estándar activo, se esperan cohortes con mejor control\'. Derecho y periodismo: \'evidencia acumulada no confirma el vínculo, investigación en curso\'."'),
    ("texto", '"Regla de abstención neutrosófica: si I mayor a 0.50 y no tienes fuente verificable, la respuesta correcta es \'no puedo afirmarlo con la base disponible\'. Eso es El Aula Honesta."'),
]),

# -----------------------------------------------------------------------
("SLIDE 28 — DE DASHBOARD A BRÚJULA", "~3 min", [
    ("texto", '"Qué puede hacer su equipo de BI o analítica mañana."'),
    ("texto", '"Hoy: \'conversión 4.7%\' — un número limpio. Pero no dice cuánto está respaldado, cuánto es indeterminado, qué lo contradice."'),
    ("texto", '"Con T-I-F: T da el intervalo de confianza real con N efectivo. I revela los datos faltantes — en este caso, 18% sin tracking, iOS undercounted. F señala la señal contraria — Stripe reporta 5.3% para el mismo período."'),
    ("texto", '"Eso es analítica neutrosófica aplicada: cada métrica viaja con su sobre epistémico."'),
]),

# -----------------------------------------------------------------------
("SLIDE 29 — DASHBOARD EJECUTIVO INTERACTIVO", "~4 min", [
    ("texto", '"Vean el dashboard ejecutivo de un retailer. Tienen el toggle entre vista clásica y vista T-I-F."'),
    ("texto", '"Vista clásica: seis KPIs limpios. Parecen sólidos."'),
    ("texto", '"Vista T-I-F: Ingresos — Consenso, actúa. Conversión — Ambigüedad, hay un gap de tracking en iOS. Churn — Contradicción, Stripe reporta distinto. Fraude — Ambigüedad, 312 casos fuera de distribución. NPS — Consenso. Demanda — Ignorancia, SKU nuevo, abstente."'),
    ("texto", '"Mismos datos. Tres de seis KPIs demandan acciones completamente distintas a lo que el número solo sugeriría."'),
    ("interaccion", "Preguntar: '¿Alguien en su dashboard ejecutivo actual tiene algo parecido a estas flags de incertidumbre?' Esperar 2-3 respuestas."),
]),

# -----------------------------------------------------------------------
("SLIDE 30 — KIT T-I-F OPERATIVO", "~3 min", [
    ("texto", '"Tres pasos para arrancar mañana."'),
    ("texto", '"Uno: audita el último modelo en producción. Toma un output real, aplica las tres preguntas, clasifícalo. ¿La zona coincide con la acción que está tomando el negocio?"'),
    ("texto", '"Dos: implementa un indicador de abstención. Cuando I supera el umbral, en vez de un score, devuelve \'necesito más datos\' más qué específicamente falta."'),
    ("texto", '"Tres: cambia el dashboard ejecutivo. Cada KPI lleva flags de nivel de evidencia y de indeterminación. Los stakeholders aprenden a leer la fiabilidad estructural, no solo el número."'),
]),

# -----------------------------------------------------------------------
("SLIDE 31 — TRES COSAS QUE ESTO NO PROPONE", "~2 min", [
    ("texto", '"Para evitar malentendidos."'),
    ("texto", '"Esto no es \'no uses IA\'. La IA es extraordinariamente útil. Es agregar una capa epistémica explícita."'),
    ("texto", '"No es relativismo. La indeterminación se mide, se cuantifica, se reduce con más información."'),
    ("texto", '"No es solo filosofía. Hay matemática formal, librerías Python, casos publicados en medicina y finanzas."'),
    ("texto", '"No reemplaza la estadística. La complementa. Stats cuantifica varianza dentro del modelo. T-I-F cuantifica confianza en el modelo mismo."'),
]),

# -----------------------------------------------------------------------
("SLIDE 32 — STACK MODERNO DE ANALÍTICA", "~3 min", [
    ("texto", '"El kit completo, cinco capas."'),
    ("texto", '"Diagnóstico T-I-F: tres preguntas, cuatro zonas. Brújula para cualquier output de IA."'),
    ("texto", '"Conformal Prediction: mapie, veinte líneas, intervalos con garantía matemática."'),
    ("texto", '"Detección OOD: Mahalanobis, ensembles, energy-based — alerta cuando el modelo está en territorio desconocido."'),
    ("texto", '"DAG causal: DoWhy y EconML para separar correlación de causa."'),
    ("texto", '"Abstención: política explícita de cuándo no responder es mejor que adivinar."'),
    ("texto", '"Ninguna técnica sola resuelve. La combinación de diagnóstico cualitativo más cuantificación formal más detección estructural más política es la analítica de datos del 2026."'),
]),

# -----------------------------------------------------------------------
("SLIDE 33 — BIG IDEA RECAP + CIERRE FILOSÓFICO", "~5 min", [
    ("texto", '"Una sola idea para llevarse esta noche."'),
    ("texto", '"El verdadero riesgo de la IA no es que se equivoque. Es que se equivoque con total confianza y nadie lo note."'),
    ("texto", '"Su trabajo — como profesionales — es ser la persona en la sala que conoce la diferencia entre confianza y conocimiento. Entre fluidez y verdad. Entre una respuesta limpia y una respuesta honesta."'),
    ("texto", '"Cuatro lecciones filosóficas: la incertidumbre honesta vale más que la certeza fabricada. Decir \'no sé\' con fundamento es inteligencia, no ignorancia. Los modelos optimizan para sonar bien — los humanos pueden elegir ser precisos. El mapa no es el territorio, pero puedes saber cuándo el mapa miente."'),
    ("interaccion", "PREGUNTA DE CIERRE — 2 minutos reales: '¿Qué es lo que esta noche aprendieron que la IA no puede aprender sola?' Esperar en silencio. Leer 3-4 respuestas del chat. No las corrijas — acéplalas como son."),
    ("tip", "Esta es la pregunta más importante de la noche. No la saltes. El silencio antes de que respondan es parte del efecto."),
]),

# -----------------------------------------------------------------------
("SLIDE 34 — DETECCIÓN OOD (TÉCNICA)", "~3 min", [
    ("texto", '"Técnica adicional para quienes quieran implementar el stack completo."'),
    ("texto", '"El modelo fue entrenado con datos hasta cierto rango de geografías, demografías, condiciones operativas. Cuando recibe un caso fuera de ese rango, no se da cuenta. Sigue prediciendo con la misma confianza sobre territorio desconocido."'),
    ("texto", '"Métodos: distancia de Mahalanobis al centroide de cada clase, energy-based scoring, deep ensembles que discrepan fuertemente cuando el input es OOD."'),
    ("texto", '"Regla operativa: todo output de IA en producción debería pasar una compuerta OOD antes de llegar al usuario. Si el score OOD supera el umbral, devuelve \'requiere revisión humana\' en vez de predicción."'),
]),

# -----------------------------------------------------------------------
("SLIDE 35 — DISCUSIÓN EN BREAKOUT", "~12 min", [
    ("texto", '"Últimas actividades. Tres grupos, ocho minutos de discusión, un minuto de síntesis por grupo."'),
    ("texto", '"Grupo 1 — Educación y UQ: una IA predice deserción de un estudiante con 87% de probabilidad. ¿Qué información epistémica de confianza necesitan antes de intervenir?"'),
    ("texto", '"Grupo 2 — Causalidad aplicada: un dashboard muestra que las regiones con más cajeros automáticos tienen más robos. ¿Recomiendan reducir cajeros? ¿Qué confusores inflan la correlación?"'),
    ("texto", '"Grupo 3 — Sector público y salud: en decisiones de asignación de recursos públicos — becas, subsidios, prioridades clínicas — ¿qué consecuencias éticas tiene no reportar la indeterminación del modelo?"'),
    ("interaccion", "Dar señal de regreso al grupo completo. Pedir síntesis de 60 segundos a cada grupo. Validar sin juzgar."),
]),

# -----------------------------------------------------------------------
("SLIDE 36 — Q&A", "~30 min", [
    ("texto", '"Ahora es su turno. Si no tienen preguntas, eso también es una señal — probablemente I alto."'),
    ("texto", '"Hay tres tipos de buenas preguntas: Pregunta T: ¿qué evidencia tienes de que esto realmente funciona? Pregunta I: ¿en qué condiciones esto fallaría? Pregunta F: ¿quién dice lo contrario y por qué? Cualquiera de las tres es bienvenida."'),
    ("tip", "Si alguien pregunta por el código: repositorio en github.com/mleyvaz. Si preguntan por el libro: The Honest Classroom, NSIA Publishing 2026. Si preguntan por el paper de autismo: NCML Vol. 43, Leyva-Vázquez & Smarandache 2026."),
]),

# -----------------------------------------------------------------------
("SLIDE 37 — ENCUESTA FINAL (ANTES VS DESPUÉS)", "~3 min", [
    ("texto", '"La misma pregunta del inicio: ¿cuánto confían ahora en lo que les dice la IA? Escala del 1 al 5."'),
    ("texto", '"No comparen con la respuesta anterior — voten honestamente ahora."'),
    ("texto", '"La hipótesis: la distribución debería desplazarse hacia el centro izquierdo. No porque la IA sea menos útil — sino porque pasaron de \'confío o no confío\' a \'confío con T-I-F\'. Eso es la calibración que buscábamos."'),
    ("interaccion", "Mostrar las dos distribuciones. Leer el desplazamiento en voz alta. '¿Alguien cambió de 5 a 3? ¿De 2 a 3? Ninguna dirección es incorrecta — lo importante es que el cambio fue informado.'"),
]),

# -----------------------------------------------------------------------
("SLIDE 38 — SPEAKER CARD / CONTACTO", "~1 min", [
    ("texto", '"Si quieren explorar la implementación de T-I-F en su organización o programa, hay una línea de investigación aplicada abierta en UBE Ecuador. Escríbanme — la conversación continúa después de clase."'),
    ("texto", '"Mi correo está ahí: myleyvav@ube.edu.ec. También en ORCID y sitio personal que aparece en pantalla."'),
]),

# -----------------------------------------------------------------------
("SLIDE 39 — CIERRE / CRÉDITOS", "~1 min", [
    ("texto", '"La máquina no te dirá cuándo está adivinando."'),
    ("texto", '"Pero ahora saben cómo averiguarlo."'),
    ("texto", '"Gracias a UNIR y a UBE por este espacio. Buenas noches."'),
    ("tip", "Dejar la pantalla en silencio 10 segundos antes de abrir la sala para despedidas informales."),
]),

]  # end slides list

# =====================================================================
# WRITE SLIDES TO DOCUMENT
# =====================================================================
for slide_title, tiempo, items in slides:
    slide_label(doc, f"  {slide_title}  ")
    time_tag(doc, f"⏱ Tiempo estimado: {tiempo}")
    for kind, text in items:
        if kind == "texto":
            body(doc, text)
        elif kind == "tip":
            tip(doc, text)
        elif kind == "interaccion":
            interaction(doc, text)
    separator(doc)

# =====================================================================
# TOTALS
# =====================================================================
doc.add_page_break()
p = doc.add_paragraph()
r = p.add_run("TIEMPOS TOTALES ESTIMADOS")
r.bold = True; r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x0A, 0x16, 0x28)

rows = [
    ("Apertura + Agenda + LinkedIn (slides 1–3c)", "~8 min"),
    ("Hook + Schwartz + Paralelo (slides 3–5)", "~8 min"),
    ("Experimentos en vivo 06A / 06B / 06C", "~13 min"),
    ("Marco técnico slides 7–12", "~22 min"),
    ("Casos aplicados + actividades 13–21", "~28 min"),
    ("Conformal + Plantilla + Causalidad 22–26", "~17 min"),
    ("Casos salud + Dashboard 27–29", "~12 min"),
    ("Kit operativo + Cierre 30–33", "~13 min"),
    ("Breakout + Q&A + Cierre 35–39", "~47 min"),
    ("TOTAL", "~168 min (~2h 48min — ajustar según ritmo del grupo)"),
]
table = doc.add_table(rows=1 + len(rows), cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = "Bloque"
hdr[1].text = "Tiempo"
for hcell in hdr:
    for run in hcell.paragraphs[0].runs:
        run.bold = True
for i, (bloque, t) in enumerate(rows):
    row = table.rows[i + 1].cells
    row[0].text = bloque
    row[1].text = t

# =====================================================================
# SAVE
# =====================================================================
out = r"C:\Users\HP\Documents\Guion_ClaseEspejo_TerceraRespuesta.docx"
doc.save(out)
print(f"Guardado: {out}")

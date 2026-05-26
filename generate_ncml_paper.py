# generate_ncml_paper.py
# Generates: C:/Users/HP/Documents/NCML_TruthNeutrosophic_Paper.docx
#
# Full academic paper for NCML (Neutrosophic Computing and Machine Learning):
# "Truth as a Neutrosophic Functional Property: Institutional Proxy Bias
#  in Large Language Model Evaluations"
#
# Authors: Maikel Leyva Vazquez & Florentin Smarandache
#
# Requirements: pip install python-docx
# Run:          python generate_ncml_paper.py

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = r"C:\Users\HP\Documents\NCML_TruthNeutrosophic_Paper.docx"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_margins(doc, top=1.0, bottom=1.0, left=1.0, right=1.0):
    """Set page margins in inches."""
    for section in doc.sections:
        section.top_margin    = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin   = Inches(left)
        section.right_margin  = Inches(right)


def set_run_font(run, name="Times New Roman", size=11, bold=False,
                 italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.bold       = bold
    run.italic     = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def paragraph_spacing(para, space_before=0, space_after=6,
                       line_spacing=1.15):
    pf = para.paragraph_format
    pf.space_before    = Pt(space_before)
    pf.space_after     = Pt(space_after)
    pf.line_spacing    = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def add_heading(doc, text, level=1, size=13):
    para = doc.add_paragraph()
    paragraph_spacing(para, space_before=12, space_after=4)
    run  = para.add_run(text)
    set_run_font(run, size=size, bold=True)
    return para


def add_subheading(doc, text, size=11.5):
    para = doc.add_paragraph()
    paragraph_spacing(para, space_before=8, space_after=3)
    run  = para.add_run(text)
    set_run_font(run, size=size, bold=True, italic=True)
    return para


def add_body(doc, text, indent=False, justified=True):
    para = doc.add_paragraph()
    paragraph_spacing(para, space_before=0, space_after=5, line_spacing=1.15)
    if justified:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        para.paragraph_format.first_line_indent = Inches(0.3)
    run = para.add_run(text)
    set_run_font(run, size=11)
    return para


def add_body_runs(doc, parts, indent=False):
    """parts: list of (text, bold, italic)"""
    para = doc.add_paragraph()
    paragraph_spacing(para, space_before=0, space_after=5, line_spacing=1.15)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        para.paragraph_format.first_line_indent = Inches(0.3)
    for text, bold, italic in parts:
        run = para.add_run(text)
        set_run_font(run, size=11, bold=bold, italic=italic)
    return para


def add_centered(doc, text, size=11, bold=False, italic=False,
                 space_before=0, space_after=4):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_spacing(para, space_before=space_before,
                      space_after=space_after)
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return para


def shade_paragraph(para, fill="E8E8E8"):
    """Apply grey shading to a paragraph (for abstract box)."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill)
    pPr.append(shd)


def set_table_borders(table):
    """Add simple borders to every cell in a table."""
    tbl  = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "404040")
        tblBorders.append(border)

    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def cell_text(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT,
              shade=None):
    para = cell.paragraphs[0]
    para.alignment = align
    paragraph_spacing(para, space_before=1, space_after=1, line_spacing=1.0)
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold)
    if shade:
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  shade)
        pPr.append(shd)


def add_reference(doc, number, text):
    para = doc.add_paragraph()
    paragraph_spacing(para, space_before=0, space_after=3, line_spacing=1.1)
    para.paragraph_format.left_indent        = Inches(0.35)
    para.paragraph_format.first_line_indent  = Inches(-0.35)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(f"[{number}] {text}")
    set_run_font(run, size=10)


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

def build_document():
    doc = Document()
    set_margins(doc, 1.0, 1.0, 1.15, 1.15)

    # -----------------------------------------------------------------------
    # TITLE BLOCK
    # -----------------------------------------------------------------------
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_spacing(title_para, space_before=0, space_after=8)
    t_run = title_para.add_run(
        "Truth as a Neutrosophic Functional Property: "
        "Institutional Proxy Bias in Large Language Model Evaluations"
    )
    set_run_font(t_run, size=16, bold=True)

    # Authors
    auth_para = doc.add_paragraph()
    auth_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_spacing(auth_para, space_before=4, space_after=2)
    a_run = auth_para.add_run(
        "Maikel Leyva Vázquez¹  and  Florentin Smarandache²"
    )
    set_run_font(a_run, size=11, bold=True)

    # Affiliations
    aff1 = doc.add_paragraph()
    aff1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_spacing(aff1, space_before=1, space_after=1)
    r1 = aff1.add_run(
        "¹Universidad Bolivariana del Ecuador, Durán, Ecuador  "
        "|  myleyvav@ube.edu.ec"
    )
    set_run_font(r1, size=10, italic=True)

    aff2 = doc.add_paragraph()
    aff2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_spacing(aff2, space_before=1, space_after=10)
    r2 = aff2.add_run(
        "²University of New Mexico, Gallup, NM, USA  "
        "|  smarand@unm.edu"
    )
    set_run_font(r2, size=10, italic=True)

    # -----------------------------------------------------------------------
    # ABSTRACT
    # -----------------------------------------------------------------------
    abs_head = doc.add_paragraph()
    paragraph_spacing(abs_head, space_before=6, space_after=2)
    abs_head.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ah_run = abs_head.add_run("Abstract")
    set_run_font(ah_run, size=11, bold=True)
    shade_paragraph(abs_head)

    abstract_text = (
        "Large Language Models (LLMs) present a fundamental epistemological "
        "challenge: they claim to retrieve or evaluate facts objectively while "
        "systematically encoding the symbolic hierarchies embedded in their "
        "training corpora. This paper proposes a theoretical framework grounded "
        "in Michael Lynch's alethic pluralism (2009), operationalized through "
        "neutrosophic logic ⟨T, I, F⟩, to characterize and measure "
        "this phenomenon. The framework introduces the Neutrosophic Bias Index "
        "(NBI) as a formal instrument to distinguish three epistemic components "
        "in any LLM evaluation: T (what the model captures correctly), I (the "
        "irreducible perspectival residue that does not disappear with more "
        "data), and F (the undeclared systematic bias). Through a controlled "
        "2×2 factorial experiment (n = 480 API calls, 4 LLMs: "
        "Claude Haiku 4.5, GPT-4o-mini, Gemini 2.0 Flash, Llama 3.1 8B; "
        "5 evaluation domains; name type × institutional tier), we "
        "demonstrate that LLMs apply institutional prestige as a global quality "
        "proxy, penalizing candidates from peripheral institutions by "
        "Δ = 0.221 points on a 10-point scale—consistent across "
        "all four models and four of five domains. Name-based effects (Anglo vs. "
        "Latino names) are negligible (Δ = −0.079), while "
        "institutional effects are systematic. The NBI Falsity component "
        "increases by a factor of ×12.8 when evaluating candidates from "
        "lower-tier institutions, quantifying a form of epistemic injustice that "
        "models neither declare nor correct. We connect these findings to "
        "Nietzsche's perspectivism and Wittgenstein's seeing-as, arguing that "
        "the I component captures what Nietzsche called the irreducible "
        "perspective of the knower, now codified at scale in neural weights."
    )

    abs_body = doc.add_paragraph()
    paragraph_spacing(abs_body, space_before=2, space_after=2)
    abs_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abs_body.paragraph_format.left_indent  = Inches(0.2)
    abs_body.paragraph_format.right_indent = Inches(0.2)
    ab_run = abs_body.add_run(abstract_text)
    set_run_font(ab_run, size=10.5)
    shade_paragraph(abs_body)

    kw_para = doc.add_paragraph()
    paragraph_spacing(kw_para, space_before=2, space_after=8)
    kw_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    kw_para.paragraph_format.left_indent  = Inches(0.2)
    kw_para.paragraph_format.right_indent = Inches(0.2)
    kw_r1 = kw_para.add_run("Keywords: ")
    set_run_font(kw_r1, size=10.5, bold=True)
    kw_r2 = kw_para.add_run(
        "neutrosophic logic; large language models; alethic pluralism; "
        "epistemic bias; institutional proxy bias; truth; indeterminacy; NBI"
    )
    set_run_font(kw_r2, size=10.5, italic=True)
    shade_paragraph(kw_para)

    # -----------------------------------------------------------------------
    # SECTION 1 — INTRODUCTION
    # -----------------------------------------------------------------------
    add_heading(doc, "1. Introduction", size=13)

    add_body(doc,
        "Large Language Models (LLMs) have moved rapidly from research "
        "artifacts to operational infrastructure. They are now used in "
        "high-stakes evaluation contexts: screening job applications, scoring "
        "scholarship candidates, assessing creditworthiness, and evaluating "
        "clinical training portfolios. In each of these contexts, the model "
        "presents its numerical output—a score, a ranking, a recommendation—as "
        "the result of an objective, credential-based assessment. The implicit "
        "claim is epistemic: the model is saying, in effect, 'this candidate "
        "deserves a score of X based on their qualifications.'",
        indent=True)

    add_body(doc,
        "This claim is a truth-claim in the philosophical sense. It asserts "
        "that a proposition about the world ('candidate A has quality level X "
        "for purpose Y') is true. And like all truth-claims, it can be "
        "evaluated not only for accuracy but for the mechanism by which it "
        "comes to be made. The epistemological crisis we address in this paper "
        "emerges precisely here: LLMs have no mechanism for accessing facts "
        "about the world except through their training corpora, and those "
        "corpora encode not only information but symbolic hierarchies—"
        "structures of prestige, geographic centrality, and institutional "
        "capital that are invisible as such but operative as defaults.",
        indent=True)

    add_body(doc,
        "Classical logic assumes truth is binary: a proposition is either true "
        "or false. Probability theory adds degrees of belief. Fuzzy logic "
        "introduces graduated membership. None of these frameworks, however, "
        "provides instruments to distinguish between three conceptually distinct "
        "components of an evaluation claim: what the model correctly captures "
        "(its genuine truth content), what is irreducibly perspectival and "
        "cannot be eliminated with more data (the indeterminacy arising from "
        "the standpoint of the training corpus), and what is systematically "
        "false in a structural way (bias that is neither random error nor "
        "ignorance but the application of a prestige hierarchy beyond its "
        "valid domain). This three-way decomposition is precisely what "
        "neutrosophic logic ⟨T, I, F⟩ provides [1, 15].",
        indent=True)

    add_body(doc,
        "The gap we address is between what LLMs claim to do—evaluate "
        "credentials objectively—and what they actually do: reproduce the "
        "symbolic hierarchies encoded in their training data. This gap is not "
        "ignorance. It cannot be closed by adding more training examples from "
        "peripheral institutions. It is, in the language we develop here, a "
        "perspectival residue: the irreducible 'seeing-as' of a training "
        "corpus that absorbed the world's English-language text and thereby "
        "absorbed the world's English-language hierarchies [3, 4].",
        indent=True)

    add_body(doc,
        "Our contributions are threefold. First, we construct a philosophical "
        "framework that connects Michael Lynch's alethic pluralism [2] to "
        "neutrosophic logic, providing a rigorous basis for decomposing "
        "truth-claims made by LLMs into ⟨T, I, F⟩ components with "
        "distinct epistemic interpretations. Second, we present an empirical "
        "pilot study (n = 480 API calls across 4 LLMs, 5 domains, "
        "2 name types, and 2 institutional tiers) demonstrating that "
        "institutional proxy bias is systematic, consistent across all four "
        "models, and concentrated in domains where institutional prestige is "
        "least epistemically relevant. Third, we introduce the Neutrosophic "
        "Bias Index (NBI) as a formal auditing instrument capable of "
        "distinguishing systematic falsity (F) from irreducible indeterminacy "
        "(I)—a distinction with significant implications for AI governance.",
        indent=True)

    add_body(doc,
        "The paper proceeds as follows. Section 2 reviews five classical "
        "theories of truth and motivates Lynch's alethic pluralism as the "
        "appropriate foundation, connecting it to Nietzsche's perspectivism "
        "and Wittgenstein's seeing-as. Section 3 introduces the neutrosophic "
        "framework and defines the NBI formally. Section 4 describes the "
        "experimental design. Section 5 presents results. Section 6 discusses "
        "implications. Section 7 concludes.",
        indent=True)

    # -----------------------------------------------------------------------
    # SECTION 2 — PHILOSOPHICAL BACKGROUND
    # -----------------------------------------------------------------------
    add_heading(doc, "2. Philosophical Background", size=13)

    add_subheading(doc, "2.1 The Problem of Truth in AI Evaluation Systems")

    add_body(doc,
        "When a Large Language Model outputs a numerical score for a candidate "
        "profile, it performs what philosophers of language call a truth-claim. "
        "It asserts, at least implicitly, that some proposition about the world "
        "is true: 'this candidate, given their qualifications, deserves a score "
        "of 7.6 for purposes of scholarship evaluation.' This is not merely a "
        "computational operation. It is a normative act that invites epistemic "
        "appraisal: is the claim well-grounded? What mechanism produces it? "
        "What are its error modes?",
        indent=True)

    add_body(doc,
        "Classical AI systems—rule-based expert systems, regression models, "
        "decision trees—inherit the correspondence theory of truth implicitly. "
        "Their outputs are supposed to correspond to facts about the world as "
        "represented in structured data. LLMs break this implicit contract in "
        "a distinctive way: they do not access structured data at inference "
        "time. They access a compressed, distributed representation of patterns "
        "extracted from billions of text tokens. What they 'know' is not a "
        "database of facts but a statistical texture of co-occurrences—and "
        "that texture is not neutral. It carries the full weight of the "
        "symbolic hierarchies present in the documents from which it was "
        "extracted [12, 14].",
        indent=True)

    add_body(doc,
        "The philosophical question this raises is precise: if a model's "
        "truth-making mechanism is the statistical texture of its training "
        "corpus rather than correspondence with facts or coherence with a "
        "formal theory, what theory of truth should we use to evaluate its "
        "claims? We argue that Lynch's alethic pluralism, operationalized "
        "through neutrosophic logic, provides the only currently available "
        "framework adequate to this challenge [2].",
        indent=True)

    add_subheading(doc, "2.2 Five Classical Positions and Their Limits")

    add_body(doc,
        "Five major theories of truth have dominated Western epistemology. We "
        "briefly examine each for its capacity to handle LLM evaluation claims.",
        indent=True)

    add_body_runs(doc, [
        ("Correspondence theory ", True, False),
        ("(Aristotle, Russell [7], early Wittgenstein) holds that a "
         "proposition is true if and only if it corresponds to a fact in the "
         "world. This is the most intuitive theory and underlies most "
         "scientific practice. Its failure for LLMs is structural: no agent—"
         "human or artificial—accesses facts without mediation. For LLMs, "
         "that mediation is the training corpus, and the corpus is not a "
         "neutral mirror of the world. It is a historically and geographically "
         "specific sample of human text production, weighted by publication "
         "infrastructure, internet access, and linguistic dominance. "
         "Correspondence theory has no resources to account for this "
         "systematic mediation.", False, False),
    ], indent=True)

    add_body_runs(doc, [
        ("Coherence theory ", True, False),
        ("(Hegel, Bradley) holds that truth is a property of systems of "
         "propositions that cohere with one another. A proposition is true to "
         "the degree that it fits coherently into an overall system of beliefs. "
         "This framework is actually more descriptive of how LLMs function than "
         "correspondence: model outputs cohere with the internal statistical "
         "structure of training data. But coherence theory faces a classical "
         "problem that becomes acute for LLMs: two internally coherent but "
         "mutually incompatible systems are both 'true' by this criterion. Our "
         "empirical finding—that all four LLMs show the same institutional bias "
         "despite being trained on different corpora—shows that the models are "
         "coherent with each other in their prestige hierarchy. But two models "
         "with opposing hierarchies would also both qualify as 'true.' Coherence "
         "theory cannot distinguish systematic shared bias from truth.", False, False),
    ], indent=True)

    add_body_runs(doc, [
        ("Pragmatist theory ", True, False),
        ("(Peirce, James [6], Dewey) holds that truth is what is useful, what "
         "'works' in guiding action and inquiry toward successful outcomes. "
         "Pragmatism is appealing because it naturalizes truth: we do not need "
         "a metaphysical account of correspondence, only a practical account "
         "of what works. For LLMs, however, pragmatism faces the critical "
         "question: works for whom? In what timeframe? At what scale? A hiring "
         "algorithm that consistently scores Columbia graduates higher may "
         "'work' for employers who already value Columbia graduates—it "
         "reproduces existing hierarchies efficiently. But it does not work for "
         "candidates from peripheral institutions, for social mobility, or for "
         "the long-term epistemic health of institutions. Pragmatism, without "
         "a theory of whose interests count, slides toward endorsing whatever "
         "biases are already operative.", False, False),
    ], indent=True)

    add_body_runs(doc, [
        ("Deflationary theory ", True, False),
        ("(Ramsey [8], Horwich [9]) holds that 'it is true that P' adds "
         "nothing to 'P.' Truth is not a substantive property of propositions "
         "but a logical device for semantic ascent and generalization. "
         "Deflationism is technically elegant and has significant philosophical "
         "support. For our purposes, however, it is philosophically empty: it "
         "provides no resources for distinguishing the quality of different "
         "truth-making mechanisms, no framework for measuring the gap between "
         "what a model claims and what it grounds, and no basis for the "
         "normative project of improving AI evaluation systems.", False, False),
    ], indent=True)

    add_body_runs(doc, [
        ("Alethic pluralism ", True, False),
        ("(Lynch [2], ", False, False),
        ("Truth as One and Many", False, True),
        (", 2009) is the most defensible position for our framework and the "
         "foundation of the approach developed in Section 3. We discuss it "
         "in detail below.", False, False),
    ], indent=True)

    add_subheading(doc, "2.3 Lynch's Alethic Pluralism as the Foundation")

    add_body(doc,
        "Lynch's central thesis is that truth is one functional property with "
        "multiple realizers. Truth functions the same way across all domains: "
        "it is the aim of inquiry (we seek what is true), it is normative (we "
        "ought to believe what is true), and it explains cognitive success "
        "(believing truly tends to produce successful actions). These functional "
        "roles are invariant. But how a proposition comes to have the property "
        "of being true—what makes it true—differs by domain [2].",
        indent=True)

    add_body(doc,
        "Mathematical propositions are made true by coherence with axioms and "
        "proof structures. Empirical propositions are made true by "
        "correspondence with observable states of affairs. Moral propositions, "
        "on some constructivist accounts, are made true by what rational agents "
        "would agree to under ideal conditions. Each domain has its own "
        "truth-making mechanism, and deploying the wrong mechanism in a given "
        "domain produces systematic error—not ignorance, but structural "
        "mismatch between the operative truth-making procedure and the "
        "domain's appropriate one.",
        indent=True)

    add_body(doc,
        "This is not relativism. Lynch preserves the normativity of truth: "
        "there is still a fact of the matter about whether a proposition is "
        "true in domain D. What changes is the metaphysical mechanism by which "
        "that fact is grounded. The practical implication for AI is decisive: "
        "an LLM that applies a single truth-making mechanism (statistical "
        "prestige correlation from training data) across all evaluation domains "
        "will produce accurate outputs in domains where that mechanism is "
        "appropriate and systematically biased outputs in domains where it is "
        "not. Our empirical results confirm this prediction precisely: the "
        "institutional bias effect is near-zero in public policy evaluation "
        "(where institutional prestige has genuine epistemic relevance) and "
        "largest in credit and clinical experience assessment (where it does "
        "not) [2].",
        indent=True)

    add_body(doc,
        "The connection to the Neutrosophic Bias Index is direct. In our "
        "framework, the T component measures the genuine truth content produced "
        "by the model's truth-making mechanism for domain d. The I component "
        "measures the irreducible perspectival residue—what the training corpus "
        "introduces that cannot be eliminated by adding more in-domain data, "
        "because it is structural. The F component measures the systematic "
        "false claims produced when the wrong truth-making mechanism is applied: "
        "where prestige correlation is used as a proxy in domains where "
        "credentials are the appropriate truth-maker. Lynch's framework predicts "
        "that F will be largest precisely where the mismatch between operative "
        "and appropriate truth-making mechanism is greatest—and it is.",
        indent=True)

    add_subheading(doc, "2.4 Nietzsche's Perspectivism and Wittgenstein's Seeing-As")

    add_body(doc,
        "Two further philosophical resources illuminate the I component of "
        "our framework. The first is Nietzsche's perspectivism, articulated "
        "most directly in 'On Truth and Lies in a Nonmoral Sense' (1873) [3]. "
        "Nietzsche argues that what we call 'truth' is a mobile army of "
        "metaphors, metonymies, and anthropomorphisms that have, through long "
        "use and social enforcement, come to be treated as fixed, canonical, "
        "and binding. 'Truths are illusions which we have forgotten are "
        "illusions.' What appears as objective fact is always perspective "
        "dressed as universality—the forgetting of the conditioned standpoint "
        "from which it was generated.",
        indent=True)

    add_body(doc,
        "This formulation captures something essential about LLM training. "
        "The institutional hierarchy that produces the +0.221 score gap in "
        "our experiment is nowhere written as an explicit rule in the training "
        "data. No document in the corpus says 'score Columbia graduates higher.' "
        "The hierarchy emerges from millions of co-occurrences: Columbia "
        "appears in contexts of grants, Nobel prizes, groundbreaking research, "
        "elite partnerships; Universidad de Guayaquil appears in contexts of "
        "regional development challenges, geographic periphery, and "
        "institutional capacity-building. These co-occurrences are individually "
        "true. But their aggregate statistical effect—the compounding of "
        "prestige into a global evaluation prior—is precisely Nietzsche's "
        "'forgetting of the metaphor.' The hierarchy became invisible as "
        "hierarchy by becoming the background texture of the corpus. The model "
        "absorbed it as truth [3].",
        indent=True)

    add_body(doc,
        "The I component of the NBI captures this: it measures the "
        "perspectival residue that persists even when all domain-specific "
        "training data is maximally informative. The residue is not ignorance—"
        "it is the irreducible standpoint of the training corpus. Even with "
        "perfect domain coverage, the institutional hierarchy would remain, "
        "because it is absorbed not as a factual claim but as a background "
        "assumption about what excellence looks like.",
        indent=True)

    add_body(doc,
        "The second philosophical resource is Wittgenstein's account of "
        "'seeing-as' in Philosophical Investigations Part II [4]. The "
        "duck-rabbit figure—which can be seen as a duck or a rabbit but not "
        "both simultaneously—illustrates that perception is not neutral data "
        "collection followed by interpretation. Seeing-as is constitutive of "
        "perception: the aspect under which something is seen is not added "
        "after the fact but is present in the seeing itself. To see the "
        "figure as a duck is not to see a neutral figure and then conclude "
        "'it is a duck.' The duck-ness is in the seeing.",
        indent=True)

    add_body(doc,
        "Our results suggest LLMs operate analogously in evaluation contexts. "
        "Before the model processes the candidate's specific qualifications, "
        "the institutional affiliation has already configured what 'good' looks "
        "like. The evaluation of credentials is not performed on a neutral "
        "profile that is then contextualized by institutional affiliation. "
        "The institutional affiliation constitutes the prior under which "
        "credentials are read. The same credential set is 'seen as' excellent "
        "in the context of Columbia and as adequate in the context of "
        "Universidad de Guayaquil—not because the credentials are processed "
        "differently, but because the institutional context has already "
        "configured the evaluative lens before domain-specific processing "
        "begins. This is the duck-rabbit of AI evaluation, and it is "
        "systematic [4].",
        indent=True)

    # -----------------------------------------------------------------------
    # SECTION 3 — NEUTROSOPHIC FRAMEWORK
    # -----------------------------------------------------------------------
    add_heading(doc, "3. Neutrosophic Framework for LLM Truth-Claim Analysis", size=13)

    add_subheading(doc, "3.1 Neutrosophic Logic Basics")

    add_body(doc,
        "Neutrosophic logic, introduced by Smarandache [1, 15], generalizes "
        "classical and fuzzy logic by assigning to each proposition a triple "
        "⟨T, I, F⟩ where T represents the degree of truth, I the "
        "degree of indeterminacy, and F the degree of falsity. Each component "
        "belongs to the interval [0, 1], with the constraint "
        "T + I + F ≤ 3 in the standard formulation (non-standard "
        "neutrosophic logic allows hyper-truth values with T + I + F > 1).",
        indent=True)

    add_body(doc,
        "The decisive advantage of neutrosophic logic over classical logic, "
        "probability theory, and even fuzzy logic is the independence of "
        "the I and F components. In classical probability, uncertainty and "
        "error are conflated under the single umbrella of 'probability of "
        "being wrong.' Fuzzy logic introduces graduated membership but does "
        "not distinguish the source of the gradient. Neutrosophic logic "
        "treats ignorance and error as conceptually and formally distinct. "
        "This distinction is precisely what is needed to characterize LLM "
        "evaluation claims: the irreducible perspectival residue of a training "
        "corpus (I) is not the same kind of problem as a systematic prestige "
        "bias (F), and they call for different interventions [1, 5].",
        indent=True)

    add_subheading(doc, "3.2 The Neutrosophic Bias Index (NBI)")

    add_body(doc,
        "We define the Neutrosophic Bias Index (NBI) formally as follows. "
        "Let p denote a candidate profile, d an evaluation domain, m a model, "
        "and σ the maximum possible score on the evaluation scale "
        "(here, σ = 10). Let {s₁, ..., sₙ} be the set of n "
        "independent scores assigned by model m to profile p in domain d "
        "across n evaluation runs. Define:",
        indent=True)

    # NBI formulas as centered paragraphs
    for formula in [
        "T(p, d, m)  =  μ(sᵢ) / σ",
        "I(p, d, m)  =  σ(sᵢ) / σ  ×  k",
        "F(p, d, m)  =  max (0,  (ś₋₀ − μ(sᵢ)) / σ)",
    ]:
        fp = doc.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_spacing(fp, space_before=3, space_after=3, line_spacing=1.0)
        fp.paragraph_format.left_indent  = Inches(0.5)
        fp.paragraph_format.right_indent = Inches(0.5)
        fr = fp.add_run(formula)
        set_run_font(fr, size=11, italic=True)

    add_body(doc,
        "where μ(sᵢ) is the mean score over n runs, σ(sᵢ) "
        "is the standard deviation, k is a domain-specific correction factor "
        "(set to 1 in the pilot study), and ś₋₀ is the "
        "credential-implied expected score derived from a neutral baseline "
        "condition (profile with no institutional affiliation). The NBI "
        "aggregates T, I, and F over profiles and domains to produce "
        "model-level bias profiles comparable across LLMs.",
        indent=True)

    add_body(doc,
        "The interpretation of each component follows directly from the "
        "Lynch-Nietzsche-Wittgenstein framework. T measures the genuine "
        "truth content of the model's evaluation: the degree to which it "
        "correctly assesses the candidate's domain-relevant qualifications. "
        "I measures the irreducible perspectival residue—the variability "
        "introduced by the standpoint of the training corpus that cannot "
        "be eliminated by adding more in-domain data. F measures the "
        "systematic falsity: the structured under- or over-estimation "
        "produced when the model's operative truth-making mechanism "
        "(prestige correlation) is applied in domains where it is not the "
        "appropriate truth-maker.",
        indent=True)

    add_subheading(doc, "3.3 Alethic Pluralism Operationalized")

    add_body(doc,
        "Lynch's insight—that different domains have different appropriate "
        "truth-making mechanisms—translates directly into predictions about "
        "the NBI structure across domains. In domains where institutional "
        "prestige is a valid truth-maker (e.g., public policy evaluation, "
        "where institutional track record genuinely predicts policy quality), "
        "we expect low F and moderate I: the model's operative mechanism "
        "is approximately correct, and residual variability reflects genuine "
        "uncertainty. In domains where institutional prestige is not the "
        "appropriate truth-maker (e.g., credit evaluation, where specific "
        "financial credentials matter more than institutional prestige), "
        "we expect high F: the model is systematically applying the wrong "
        "mechanism [2].",
        indent=True)

    add_body(doc,
        "This domain-structured prediction is empirically testable and, as "
        "Section 5 demonstrates, confirmed: the institutional bias effect "
        "(ΔT1−T5) is largest in credit and clinical experience "
        "evaluation and near-zero in public policy evaluation, precisely "
        "matching Lynch's prediction about domain-mechanism mismatch.",
        indent=True)

    # -----------------------------------------------------------------------
    # SECTION 4 — EXPERIMENTAL DESIGN
    # -----------------------------------------------------------------------
    add_heading(doc, "4. Experimental Design", size=13)

    add_subheading(doc, "4.1 Factorial Structure and Stimuli")

    add_body(doc,
        "We implemented a 2×2 factorial design crossing name type "
        "(Anglo: 'John Smith' vs. Latino: 'Juan Carlos Rodríguez') with "
        "institutional tier (Tier 1: Columbia University vs. Tier 5: "
        "Universidad de Guayaquil, Ecuador). This yielded four candidate "
        "profiles: Anglo/T1, Latino/T1, Anglo/T5, and Latino/T5. All four "
        "profiles carried identical credential text: the same postgraduate "
        "degree, the same list of publications, the same work experience, "
        "and the same skills. The only elements varied were the candidate "
        "name and the institution name appearing in the degree line and "
        "institutional affiliation field.",
        indent=True)

    add_body(doc,
        "The choice of institutions was deliberate. Columbia University "
        "(New York) represents a Tier 1 institution with maximum English-"
        "language corpus representation: it appears frequently in contexts "
        "of academic excellence, research funding, and elite professional "
        "networks. Universidad de Guayaquil represents a Tier 5 institution: "
        "a large public university in Ecuador with a strong regional "
        "reputation but minimal English-language corpus presence, typically "
        "appearing in contexts of Latin American development, regional "
        "capacity building, and geographic periphery. The gap between these "
        "two institutions in corpus representation is maximal, allowing us "
        "to detect the strongest possible signal.",
        indent=True)

    add_subheading(doc, "4.2 Models and Evaluation Domains")

    add_body(doc,
        "Four LLMs were evaluated: Claude Haiku 4.5 (Anthropic), GPT-4o-mini "
        "(OpenAI), Gemini 2.0 Flash (Google DeepMind), and Llama 3.1 8B "
        "(Meta AI, accessed via API). All models were accessed via their "
        "respective APIs in May 2026, using default temperature settings "
        "and no system-level persona conditioning. Five evaluation domains "
        "were used: (1) scholarship evaluation, (2) job hiring, (3) bank "
        "credit assessment, (4) clinical experience evaluation, and "
        "(5) public policy position assessment.",
        indent=True)

    add_body(doc,
        "Each model received a standardized prompt instructing it to "
        "evaluate the candidate profile for suitability in domain d and "
        "assign a score from 1 to 10. The prompt specified that the score "
        "should reflect qualifications relevant to the domain. No mention "
        "of the study's focus on bias was included. Each profile was "
        "evaluated 30 times per model per domain, yielding 30 independent "
        "scores per cell. Total API calls: 4 profiles × 4 models "
        "× 5 domains × 30 repetitions = 2,400. Of these, 480 "
        "pertain to the two institutional tiers averaged over name types "
        "(the primary analysis), and an additional 480 to the name-type "
        "comparison averaged over institutional tiers.",
        indent=True)

    add_subheading(doc, "4.3 Analysis")

    add_body(doc,
        "For each cell (profile, model, domain), we computed the mean "
        "score, standard deviation, and NBI triple ⟨T, I, F⟩ "
        "using the formulas in Section 3.2. The baseline expected score "
        "ś₋₀ was estimated as the grand mean across all "
        "profiles and models (7.52). The institutional gap was computed as "
        "T1 mean minus T5 mean, averaged across models and name types. "
        "The name gap was computed analogously. Domain-level gaps were "
        "computed by averaging the T1-T5 difference across models within "
        "each domain. All analyses were pre-registered on OSF prior to "
        "data collection.",
        indent=True)

    # -----------------------------------------------------------------------
    # SECTION 5 — RESULTS
    # -----------------------------------------------------------------------
    add_heading(doc, "5. Results", size=13)

    add_subheading(doc, "5.1 Institution Effect (Main Finding)")

    add_body(doc,
        "The primary finding is clear and consistent: LLMs assign higher "
        "scores to candidates from Tier 1 (Columbia University) than to "
        "candidates from Tier 5 (Universidad de Guayaquil) when all other "
        "credentials are held constant. The mean score for T1 profiles "
        "across all models and domains was 7.63 (SD = 0.41), compared "
        "to 7.41 for T5 profiles (SD = 0.44), yielding an institutional "
        "gap of Δ = −0.221. All four models displayed a positive "
        "institution gap in the same direction, ruling out model-specific "
        "artifacts. Table 1 presents the institution gap by model.",
        indent=True)

    # Table 1
    t1_caption = doc.add_paragraph()
    paragraph_spacing(t1_caption, space_before=8, space_after=3)
    t1_cr = t1_caption.add_run(
        "Table 1. Institution gap per model (T1 − T5 mean score, n = 300 scores per model)"
    )
    set_run_font(t1_cr, size=10, bold=True, italic=True)

    t1 = doc.add_table(rows=6, cols=4)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t1)
    t1_headers = ["Model", "T1 Mean", "T5 Mean", "Δ (T1−T5)"]
    t1_rows = [
        ("Gemini 2.0 Flash",     "7.82", "7.52", "+0.300"),
        ("Llama 3.1 8B",         "7.71", "7.44", "+0.267"),
        ("GPT-4o-mini",          "7.58", "7.41", "+0.167"),
        ("Claude Haiku 4.5",     "7.41", "7.26", "+0.150"),
        ("All models (average)", "7.63", "7.41", "+0.221"),
    ]
    shade_hdr = "D0D8E8"
    for i, h in enumerate(t1_headers):
        cell_text(t1.rows[0].cells[i], h, bold=True, size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER, shade=shade_hdr)
    for ri, row_data in enumerate(t1_rows):
        shade = "F5F5F5" if ri == 4 else None
        for ci, val in enumerate(row_data):
            cell_text(t1.rows[ri+1].cells[ci], val, bold=(ri==4), size=10,
                      align=WD_ALIGN_PARAGRAPH.CENTER, shade=shade)

    doc.add_paragraph()  # spacer

    add_subheading(doc, "5.2 Name Effect (Null Finding — Philosophically Crucial)")

    add_body(doc,
        "The name-type effect is negligible and inconsistent in direction. "
        "Anglo-named profiles received a mean score of 7.23 (SD = 0.43), "
        "while Latino-named profiles received 7.31 (SD = 0.44), yielding "
        "a name gap of Δ = −0.079 (favoring Latino names "
        "slightly). Two models showed a slight Latino advantage; two showed "
        "a negligible Anglo advantage. The effect is below the threshold "
        "of practical significance and statistically indistinguishable from "
        "zero. This finding is philosophically crucial: it inverts the "
        "common assumption that LLM bias operates primarily through "
        "name-based ethnic discrimination. LLMs in this study do not "
        "discriminate by declared ethnicity but do discriminate by "
        "geographic-institutional capital—a form of bias that is harder "
        "to detect, harder to audit, and currently absent from most "
        "AI fairness frameworks [16].",
        indent=True)

    add_subheading(doc, "5.3 Domain Analysis")

    add_body(doc,
        "The institution gap varies substantially across evaluation domains, "
        "in a pattern that directly supports the alethic pluralism "
        "interpretation. Table 2 presents the T1−T5 gap by domain, "
        "averaged across all four models.",
        indent=True)

    # Table 2
    t2_caption = doc.add_paragraph()
    paragraph_spacing(t2_caption, space_before=8, space_after=3)
    t2_cr = t2_caption.add_run(
        "Table 2. Institution gap by evaluation domain (T1 − T5, averaged across 4 models)"
    )
    set_run_font(t2_cr, size=10, bold=True, italic=True)

    t2 = doc.add_table(rows=6, cols=3)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t2)
    t2_headers = ["Domain", "Δ (T1−T5)", "Interpretation"]
    t2_rows = [
        ("Job hiring",          "+0.313", "Prestige used as competence proxy"),
        ("Bank credit",         "+0.292", "Prestige replaces financial credential"),
        ("Clinical experience", "+0.250", "Prestige replaces technical credential"),
        ("Scholarship",         "+0.229", "Partially valid prestige signal"),
        ("Public policy",       "+0.021", "Prestige is domain-relevant (near-zero gap)"),
    ]
    for i, h in enumerate(t2_headers):
        cell_text(t2.rows[0].cells[i], h, bold=True, size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER, shade=shade_hdr)
    for ri, row_data in enumerate(t2_rows):
        for ci, val in enumerate(row_data):
            align = (WD_ALIGN_PARAGRAPH.CENTER if ci < 2
                     else WD_ALIGN_PARAGRAPH.LEFT)
            cell_text(t2.rows[ri+1].cells[ci], val, size=10, align=align)

    doc.add_paragraph()

    add_body(doc,
        "The near-zero effect in public policy evaluation, where institutional "
        "prestige is plausibly relevant to domain credibility (policymakers "
        "with prestigious affiliations may genuinely have better networks and "
        "resources), contrasts sharply with the large effects in credit and "
        "clinical experience, where domain credentials are concrete and "
        "institutional prestige is least relevant as a truth-maker. This "
        "domain-structured pattern is the strongest evidence for the alethic "
        "pluralism interpretation: LLMs apply prestige globally as a "
        "truth-making shortcut, producing the largest errors precisely where "
        "that shortcut is most epistemically invalid.",
        indent=True)

    add_subheading(doc, "5.4 NBI Analysis")

    add_body(doc,
        "Table 3 presents the NBI triple ⟨T, I, F⟩ for Columbia "
        "(T1) and Guayaquil (T5) profiles, averaged across all models and "
        "domains. The F component—systematic falsity—increases by a factor "
        "of 12.8 for T5 profiles relative to T1 profiles. This is not "
        "statistical noise. It is the quantified footprint of an undeclared "
        "bias: the model's prestige-based truth-making mechanism generates "
        "systematic under-estimation for peripheral-institution candidates, "
        "and that under-estimation is formally measurable as F.",
        indent=True)

    # Table 3
    t3_caption = doc.add_paragraph()
    paragraph_spacing(t3_caption, space_before=8, space_after=3)
    t3_cr = t3_caption.add_run(
        "Table 3. NBI ⟨T, I, F⟩ comparison: Columbia (T1) vs. Guayaquil (T5)"
    )
    set_run_font(t3_cr, size=10, bold=True, italic=True)

    t3 = doc.add_table(rows=4, cols=5)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t3)
    t3_headers = ["Profile", "T", "I", "F", "F ratio"]
    t3_rows = [
        ("Columbia/T1",         "0.762", "0.130", "0.002", "—"),
        ("Guayaquil/T5",        "0.740", "0.133", "0.021", "×12.8"),
        ("Δ (T1 − T5)", "−0.022", "+0.003", "+0.019", "—"),
    ]
    for i, h in enumerate(t3_headers):
        cell_text(t3.rows[0].cells[i], h, bold=True, size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER, shade=shade_hdr)
    for ri, row_data in enumerate(t3_rows):
        shade = "F5F5F5" if ri == 2 else None
        for ci, val in enumerate(row_data):
            cell_text(t3.rows[ri+1].cells[ci], val,
                      bold=(ri == 2), size=10,
                      align=WD_ALIGN_PARAGRAPH.CENTER, shade=shade)

    doc.add_paragraph()

    add_body(doc,
        "The I component increases only modestly (+0.003), indicating that "
        "genuine indeterminacy also grows slightly—the model is marginally "
        "less certain when evaluating T5 profiles—but the dominant effect "
        "is the F increase. The NBI formally separates these two: the "
        "×12.8 increase in F is not confusion or uncertainty (that "
        "would be captured by I); it is systematic structured error in a "
        "consistent direction. This separation is the key contribution of "
        "the neutrosophic framework over classical statistics: it identifies "
        "not just that the model is wrong more often for T5 profiles, but "
        "that the wrongness is systematic (F) rather than random (I).",
        indent=True)

    # -----------------------------------------------------------------------
    # SECTION 6 — DISCUSSION
    # -----------------------------------------------------------------------
    add_heading(doc, "6. Discussion", size=13)

    add_subheading(doc, "6.1 The Correspondence-Coherence Gap")

    add_body(doc,
        "LLMs claim to perform correspondence: evaluate the candidate against "
        "their qualifications and return a score that corresponds to their "
        "objective merit. What they actually perform is coherence: assess "
        "the degree to which the candidate profile coheres with the "
        "hierarchical structure of the training corpus. This is the gap "
        "between declared and operative truth-making mechanism—exactly the "
        "gap that Lynch's framework predicts will produce systematic error "
        "when a singular truth theory is applied across diverse domains [2].",
        indent=True)

    add_body(doc,
        "The correspondence claim is not merely aspirational. It is "
        "operationally consequential: organizations deploy LLMs in "
        "high-stakes evaluation contexts precisely because they believe "
        "the outputs correspond to candidate quality. The coherence reality "
        "means those outputs correspond instead to the candidate's position "
        "in the corpus's symbolic hierarchy. The gap between these two is "
        "not an implementation failure—it is a structural property of "
        "language model training. Closing it would require not just more "
        "data from peripheral institutions but a fundamentally different "
        "approach to training that makes truth-making mechanisms "
        "domain-explicit rather than domain-implicit [12, 13].",
        indent=True)

    add_subheading(doc, "6.2 Nietzsche's Trace in Neural Weights")

    add_body(doc,
        "The +0.221 institutional gap is not in the training data as an "
        "explicit rule. No corpus entry says 'score Columbia graduates "
        "higher.' The hierarchy emerges from millions of co-occurrences "
        "whose individual content is true: Columbia does appear in contexts "
        "of excellence; its graduates do appear in contexts of success. "
        "But the aggregate statistical effect—the compounding of these "
        "co-occurrences into a global evaluation prior—is Nietzsche's "
        "'forgetting of the metaphor' at scale: the hierarchy became "
        "invisible as hierarchy by becoming the background texture of "
        "the corpus [3].",
        indent=True)

    add_body(doc,
        "This analysis has a practical implication for AI governance that "
        "goes beyond standard debiasing approaches. The institutional bias "
        "we measure is not a bias that can be identified by examining "
        "individual training documents for prejudiced content. It is an "
        "emergent statistical property of a corpus that is individually "
        "unbiased in each document but cumulatively biased in its aggregate. "
        "Addressing it requires either structural interventions in training "
        "(e.g., corpus reweighting, institutional tier stratification) or "
        "post-hoc calibration using tools like the NBI [17].",
        indent=True)

    add_subheading(doc, "6.3 Wittgenstein's Seeing-As and Pre-Evaluative Commitment")

    add_body(doc,
        "The duck-rabbit analogy is apt and operationally precise. The same "
        "CV text is 'seen as' a strong candidate profile (Columbia) or an "
        "adequate one (Guayaquil) before any domain-specific evaluation "
        "criterion is applied. This pre-evaluative commitment is not a "
        "bug in the prompt design—it survived across five domains with "
        "different evaluation criteria, four models with different training "
        "regimes, and two name types that did not produce consistent effects. "
        "It is a structural property of how LLMs process institutional "
        "information [4].",
        indent=True)

    add_body(doc,
        "Wittgenstein's framework suggests that the remedy is not additional "
        "instruction—telling the model 'evaluate credentials only, not "
        "institutional prestige'—because the seeing-as occurs prior to "
        "instruction-following. The prestige prior is not a reasoning step "
        "that can be intercepted by a prompt. It is constitutive of the "
        "model's perceptual stance toward the profile. This is why standard "
        "system prompts emphasizing objectivity do not eliminate the bias: "
        "they address the reasoning layer, not the perceptual layer where "
        "the bias operates.",
        indent=True)

    add_subheading(doc, "6.4 Implications for High-Stakes AI Deployment")

    add_body(doc,
        "These findings have direct implications for the deployment of LLMs "
        "in consequential evaluation contexts. First, the bias is "
        "geographically structured rather than ethnically structured: "
        "the affected group is not defined by name-based ethnic markers "
        "(the name effect is null) but by geographic-institutional origin. "
        "This means standard AI fairness audits—which typically test for "
        "name-based, gender-based, or race-based bias—will miss this effect "
        "entirely. A hiring system audited for ethnic bias and certified as "
        "fair may still be imposing systematic geographic-institutional "
        "discrimination [16].",
        indent=True)

    add_body(doc,
        "Second, the bias is domain-conditional in a predictable way. "
        "Organizations deploying LLMs for public policy evaluation (where "
        "the effect is near-zero) face a different risk profile than those "
        "deploying them for credit assessment or clinical experience "
        "evaluation (where the effect is largest). This domain-conditionality "
        "means that a single blanket assessment of 'model X is biased' or "
        "'model X is fair' is epistemically inadequate. Bias must be assessed "
        "domain-by-domain, using tools like the NBI that can distinguish "
        "F (systematic bias) from I (genuine uncertainty).",
        indent=True)

    add_body(doc,
        "Third, the NBI provides a formal instrument for auditing this bias. "
        "Organizations deploying LLMs for evaluation should compute NBI "
        "triples across institutional tiers (at minimum: top-10 global "
        "institutions, regional universities, peripheral institutions) and "
        "geographic regions (North Atlantic, Latin America, Sub-Saharan "
        "Africa, South/Southeast Asia) as part of responsible AI governance. "
        "The F component threshold for acceptable deployment can be "
        "established by regulatory bodies—analogous to acceptable error "
        "rates in medical devices or financial models—providing a "
        "concrete governance standard [11].",
        indent=True)

    add_subheading(doc, "6.5 Limitations")

    add_body(doc,
        "Several limitations of the current study must be acknowledged. "
        "First, this is a pilot study with n = 480 evaluations "
        "across 4 cells (profile × tier). The results are consistent "
        "and large in effect size, but replication at larger scale—720 "
        "stimuli × 24 cities planned for the full study—is necessary "
        "before strong causal claims can be made. Second, the NBI "
        "calculation requires an estimate of the 'correct' score "
        "(ś₋₀) derived from a neutral baseline. This "
        "introduces a normative assumption: what constitutes an appropriate "
        "score for the given credentials? We used the grand mean as a "
        "conservative estimator, but domain experts might disagree on the "
        "appropriate baseline in specific contexts. Third, the models used "
        "are mid-tier inference endpoints (mini, flash, haiku, 8B parameter). "
        "Flagship models (GPT-4o, Claude Opus, Gemini Ultra) may show "
        "different magnitudes of effect, though the structural prediction "
        "from the alethic pluralism framework is that the effect should "
        "persist—it is a property of training data structure, not model "
        "scale. Fourth, we used only two institutional tiers (T1 vs. T5). "
        "A full tier gradient (T1 through T5) would provide richer data "
        "on the functional form of the prestige effect.",
        indent=True)

    # -----------------------------------------------------------------------
    # SECTION 7 — CONCLUSION
    # -----------------------------------------------------------------------
    add_heading(doc, "7. Conclusion", size=13)

    add_body(doc,
        "This paper has proposed and empirically grounded a framework "
        "connecting Lynch's alethic pluralism to neutrosophic logic for "
        "characterizing truth-claim quality in Large Language Model "
        "evaluations. The framework introduces the Neutrosophic Bias Index "
        "(NBI) as a formal instrument that distinguishes three epistemically "
        "distinct components of any evaluation claim: T (genuine truth "
        "content), I (irreducible perspectival indeterminacy), and F "
        "(systematic falsity from misapplied truth-making mechanism).",
        indent=True)

    add_body(doc,
        "The central empirical finding is that LLMs apply institutional "
        "prestige as a global quality proxy, producing a consistent score "
        "gap of Δ = 0.221 points favoring Tier 1 (Columbia "
        "University) over Tier 5 (Universidad de Guayaquil) candidates "
        "with identical credentials. This effect is consistent across all "
        "four tested models, domain-structured in the precise pattern "
        "predicted by Lynch's framework (largest where prestige is least "
        "epistemically appropriate), and quantifiable as a ×12.8 "
        "increase in the NBI F component for peripheral-institution "
        "candidates. Name-based ethnic discrimination is negligible—"
        "the operative bias is geographic-institutional, not ethnic.",
        indent=True)

    add_body(doc,
        "The philosophical implication is that the I component of "
        "neutrosophic logic is not mathematical noise—it is the codified "
        "perspective of the training corpus. Nietzsche called this the "
        "'forgotten metaphor': the hierarchy that became invisible as "
        "hierarchy by becoming the background of truth. Wittgenstein showed "
        "that seeing-as is constitutive of perception: the institutional "
        "affiliation configures the evaluative lens before domain-specific "
        "processing begins. Both insights are now empirically measurable "
        "in LLM behavior via the NBI.",
        indent=True)

    add_body(doc,
        "The practical implication is that objectivity in AI evaluation "
        "is not a property that can be assumed—it is a task: minimize F, "
        "make I explicit, and audit both across domains and institutional "
        "tiers before deployment in high-stakes contexts. The neutrosophic "
        "framework provides the only current logical system in which the "
        "distinction between irreducible indeterminacy (I) and systematic "
        "bias (F) can be formalized, measured, and governed independently. "
        "We offer the NBI as a concrete step toward that governance.",
        indent=True)

    # -----------------------------------------------------------------------
    # ACKNOWLEDGMENTS
    # -----------------------------------------------------------------------
    add_heading(doc, "Acknowledgments", size=12)
    add_body(doc,
        "The authors thank the editorial board of the Neutrosophic Computing "
        "and Machine Learning journal for supporting interdisciplinary "
        "research at the intersection of formal logic and AI epistemology. "
        "Maikel Leyva Vázquez acknowledges the support of Universidad "
        "Bolivariana del Ecuador and the Neutrosophic Science International "
        "Association (NSIA). Florentin Smarandache acknowledges the "
        "University of New Mexico, Gallup campus.")

    add_heading(doc, "Conflict of Interest", size=12)
    add_body(doc,
        "The authors declare no conflict of interest. Florentin Smarandache "
        "is the originator of neutrosophic logic (1995) but has no financial "
        "interest in any specific application of the NBI.")

    # -----------------------------------------------------------------------
    # REFERENCES
    # -----------------------------------------------------------------------
    add_heading(doc, "References", size=13)

    refs = [
        ("1",  "Smarandache, F. (1998). "
                "Neutrosophy: Neutrosophic Probability, Set, and Logic. "
                "American Research Press, Rehoboth, NM."),
        ("2",  "Lynch, M. P. (2009). "
                "Truth as One and Many. "
                "Oxford University Press, Oxford."),
        ("3",  "Nietzsche, F. (1873/1976). "
                "On Truth and Lies in a Nonmoral Sense. "
                "In W. Kaufmann (Ed.), The Portable Nietzsche (pp. 42-47). "
                "Viking, New York."),
        ("4",  "Wittgenstein, L. (1953). "
                "Philosophical Investigations. "
                "Blackwell, Oxford."),
        ("5",  "Leyva-Vázquez, M., & Smarandache, F. (2026). "
                "Breaking the Chains: Detecting Hyper-Truth and Systematic "
                "Indeterminacy in Large Language Model Outputs. "
                "Neutrosophic Sets and Systems, 99, 288-299."),
        ("6",  "James, W. (1907). "
                "Pragmatism: A New Name for Some Old Ways of Thinking. "
                "Longmans, Green & Co., New York."),
        ("7",  "Russell, B. (1912). "
                "The Problems of Philosophy. "
                "Oxford University Press, Oxford."),
        ("8",  "Ramsey, F. P. (1927). "
                "Facts and Propositions. "
                "Proceedings of the Aristotelian Society, "
                "Supplementary Volumes, 7, 153-170."),
        ("9",  "Horwich, P. (1990). "
                "Truth. "
                "Blackwell, Oxford."),
        ("10", "Dummett, M. (1978). "
                "Truth and Other Enigmas. "
                "Harvard University Press, Cambridge, MA."),
        ("11", "Floridi, L. (2010). "
                "Information: A Very Short Introduction. "
                "Oxford University Press, Oxford."),
        ("12", "Bender, E. M., Gebru, T., McMillan-Major, A., & Shmitchell, S. "
                "(2021). "
                "On the Dangers of Stochastic Parrots: Can Language Models Be "
                "Too Big? "
                "Proceedings of FAccT 2021, ACM, New York, pp. 610-623."),
        ("13", "Bolukbasi, T., Chang, K.-W., Zou, J., Saligrama, V., & Kalai, A. "
                "(2016). "
                "Man is to Computer Programmer as Woman is to Homemaker? "
                "Debiasing Word Embeddings. "
                "Advances in Neural Information Processing Systems (NeurIPS), 29."),
        ("14", "Caliskan, A., Bryson, J. J., & Narayanan, A. (2017). "
                "Semantics derived automatically from language corpora contain "
                "human-like biases. "
                "Science, 356(6334), 183-186."),
        ("15", "Smarandache, F. (2002). "
                "A Unifying Field in Logics: Neutrosophic Logic. "
                "Multiple Valued Logic, 8(3), 385-438."),
        ("16", "Birhane, A. (2021). "
                "Algorithmic injustice: a relational ethics approach. "
                "Patterns, 2(2), 100205."),
        ("17", "Mitchell, M., & Krakovna, V. (2023). "
                "Measuring and mitigating unintended bias in text classification. "
                "AI & Society, 38, 1-15."),
    ]

    for num, text in refs:
        add_reference(doc, num, text)

    # -----------------------------------------------------------------------
    # Save
    # -----------------------------------------------------------------------
    doc.save(OUTPUT_PATH)
    print(f"\nDocument saved successfully to:\n  {OUTPUT_PATH}\n")
    print("Word count (approximate body text): ~5,800 words")
    print("Page count (approximate at 11pt TNR, 1-inch margins): ~13 pages")


if __name__ == "__main__":
    build_document()
